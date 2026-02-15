#!/usr/bin/env python3
"""Knowledge-base indexing and retrieval for OpenClaw V4.1."""

from __future__ import annotations

import csv
import json
import os
import re
import sqlite3
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document

from scripts.skill_clawrag_bridge import clawrag_delete, clawrag_search, clawrag_sync
from scripts.v4_runtime import (
    KB_SUPPORTED_EXTENSIONS,
    SOURCE_GROUP_WEIGHTS,
    compute_sha256,
    infer_source_group,
    json_dumps,
    resolve_rag_collection,
    utc_now_iso,
)

_KB_FTS_AVAILABLE: bool | None = None

_RERANK_SOURCE_GROUP_WEIGHTS: dict[str, float] = {
    "glossary": 2.2,
    "previously_translated": 1.6,
    "translated_en": 1.3,
    "arabic_source": 1.0,
    "general": 0.9,
}


def _rank_semantic(rank: int | None) -> float:
    if rank is None:
        return 0.0
    try:
        idx = max(0, int(rank))
    except (TypeError, ValueError):
        return 0.0
    return 1.0 / (1.0 + float(idx))


def _task_boost(task_type: str, source_group: str) -> float:
    task = (task_type or "").upper().strip()
    boosts = {
        "REVISION_UPDATE": {"glossary": 1.25, "previously_translated": 1.2},
        "NEW_TRANSLATION": {"glossary": 1.2, "arabic_source": 1.1},
        "BILINGUAL_REVIEW": {"glossary": 1.2, "translated_en": 1.15},
        "EN_ONLY_EDIT": {"translated_en": 1.2, "previously_translated": 1.1},
        "MULTI_FILE_BATCH": {"glossary": 1.15, "previously_translated": 1.1},
        "TERMINOLOGY_ENFORCEMENT": {"glossary": 1.4, "translated_en": 1.15},
        "FORMAT_CRITICAL_TASK": {"glossary": 1.2, "previously_translated": 1.15},
        "LOW_CONTEXT_TASK": {"glossary": 1.1},
    }.get(task, {})
    group = (source_group or "general").strip() or "general"
    return float(boosts.get(group, 1.0))


def _compute_rerank_score(*, semantic_score: float, source_group: str, task_type: str) -> dict[str, float]:
    sg = (source_group or "general").strip() or "general"
    sg_weight = float(_RERANK_SOURCE_GROUP_WEIGHTS.get(sg, _RERANK_SOURCE_GROUP_WEIGHTS["general"]))
    boost = _task_boost(task_type, sg)
    # Unified formula (default initial weights)
    final_score = 0.45 * float(semantic_score) + 0.35 * sg_weight + 0.20 * boost
    return {"source_group_weight": sg_weight, "task_boost": boost, "final_score": float(final_score)}


def _merge_and_rerank_hits(
    *,
    rag_hits: list[dict[str, Any]],
    local_hits: list[dict[str, Any]],
    task_type: str,
    final_k: int,
    glossary_min: int,
    terminology_glossary_ratio: float,
    prefer_rag_ratio: float,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Merge (rag + local) candidates and re-rank with hard constraints and soft quotas."""
    candidates_by_key: dict[tuple[str, int], dict[str, Any]] = {}

    for idx, hit in enumerate(rag_hits):
        path = str(hit.get("path") or "").strip()
        chunk = int(hit.get("chunk_index") or 0)
        key = (path, chunk)
        if not path:
            continue
        sg = str(hit.get("source_group") or "general").strip() or "general"
        candidates_by_key[key] = {
            "path": path,
            "chunk_index": chunk,
            "snippet": str(hit.get("snippet") or ""),
            "source_group": sg,
            "rag_rank": idx,
            "local_rank": None,
            "origin": "clawrag",
        }

    for idx, hit in enumerate(local_hits):
        path = str(hit.get("path") or "").strip()
        chunk = int(hit.get("chunk_index") or 0)
        key = (path, chunk)
        if not path:
            continue
        sg = str(hit.get("source_group") or "general").strip() or "general"
        existing = candidates_by_key.get(key)
        if existing:
            existing["local_rank"] = idx
            # Prefer local-inferred source_group when available.
            existing["source_group"] = sg or existing.get("source_group") or "general"
            if not existing.get("snippet") and hit.get("snippet"):
                existing["snippet"] = str(hit.get("snippet") or "")
            existing["origin"] = "both"
        else:
            candidates_by_key[key] = {
                "path": path,
                "chunk_index": chunk,
                "snippet": str(hit.get("snippet") or ""),
                "source_group": sg,
                "rag_rank": None,
                "local_rank": idx,
                "origin": "local",
            }

    candidates: list[dict[str, Any]] = []
    for cand in candidates_by_key.values():
        rag_rank = cand.get("rag_rank")
        local_rank = cand.get("local_rank")
        semantic = max(_rank_semantic(rag_rank), _rank_semantic(local_rank))
        sg = str(cand.get("source_group") or "general").strip() or "general"
        scored = _compute_rerank_score(semantic_score=semantic, source_group=sg, task_type=task_type)
        cand["semantic_score"] = float(semantic)
        cand["source_group_weight"] = float(scored["source_group_weight"])
        cand["task_boost"] = float(scored["task_boost"])
        cand["final_score"] = float(scored["final_score"])
        candidates.append(cand)

    candidates.sort(
        key=lambda x: (
            float(x.get("final_score") or 0.0),
            float(x.get("semantic_score") or 0.0),
            float(x.get("source_group_weight") or 0.0),
            str(x.get("path") or ""),
            int(x.get("chunk_index") or 0),
        ),
        reverse=True,
    )

    final_k = max(1, int(final_k))
    rag_target = int(round(float(prefer_rag_ratio) * float(final_k)))
    rag_target = max(0, min(final_k, rag_target))
    local_target = final_k - rag_target

    glossary_candidates = [c for c in candidates if str(c.get("source_group") or "") == "glossary"]
    has_glossary = bool(glossary_candidates)
    forced_glossary = 0

    glossary_needed = 0
    task_upper = (task_type or "").upper().strip()
    if has_glossary:
        glossary_needed = max(0, int(glossary_min))
        if task_upper == "TERMINOLOGY_ENFORCEMENT":
            ratio_needed = int((float(terminology_glossary_ratio) * float(final_k) + 0.999999))
            glossary_needed = max(glossary_needed, ratio_needed)
        glossary_needed = min(glossary_needed, len(glossary_candidates))

    selected: list[dict[str, Any]] = []
    selected_keys: set[tuple[str, int]] = set()
    selected_by_backend = {"clawrag": 0, "local": 0}

    def _supports_backend(c: dict[str, Any], backend: str) -> bool:
        if backend == "clawrag":
            return c.get("rag_rank") is not None
        if backend == "local":
            return c.get("local_rank") is not None
        return False

    def _assign_backend(c: dict[str, Any], preferred: str) -> str:
        if c.get("origin") == "clawrag":
            return "clawrag"
        if c.get("origin") == "local":
            return "local"
        # both: pick preferred when possible, otherwise the other.
        if preferred in {"clawrag", "local"} and _supports_backend(c, preferred):
            return preferred
        return "local" if preferred == "clawrag" else "clawrag"

    # Stage 1: forced glossary (hard rule)
    if glossary_needed:
        glossary_sorted = sorted(
            glossary_candidates,
            key=lambda x: (
                float(x.get("final_score") or 0.0),
                float(x.get("semantic_score") or 0.0),
            ),
            reverse=True,
        )
        for cand in glossary_sorted[:glossary_needed]:
            key = (str(cand.get("path") or ""), int(cand.get("chunk_index") or 0))
            if key in selected_keys:
                continue
            preferred = "clawrag" if selected_by_backend["clawrag"] < rag_target else "local"
            assigned = _assign_backend(cand, preferred)
            selected_by_backend[assigned] += 1
            selected.append(cand)
            selected_keys.add(key)
            forced_glossary += 1

    # Stage 2: fill by score with soft quotas
    while len(selected) < final_k:
        rag_need = rag_target - selected_by_backend["clawrag"]
        local_need = local_target - selected_by_backend["local"]
        if rag_need > local_need:
            preferred = "clawrag"
        elif local_need > rag_need:
            preferred = "local"
        else:
            preferred = "clawrag" if rag_need > 0 else "any"

        pick = None
        if preferred in {"clawrag", "local"}:
            for cand in candidates:
                key = (str(cand.get("path") or ""), int(cand.get("chunk_index") or 0))
                if key in selected_keys:
                    continue
                if _supports_backend(cand, preferred):
                    pick = cand
                    break
            if pick is None:
                other = "local" if preferred == "clawrag" else "clawrag"
                for cand in candidates:
                    key = (str(cand.get("path") or ""), int(cand.get("chunk_index") or 0))
                    if key in selected_keys:
                        continue
                    if _supports_backend(cand, other):
                        pick = cand
                        preferred = other
                        break
        if pick is None:
            for cand in candidates:
                key = (str(cand.get("path") or ""), int(cand.get("chunk_index") or 0))
                if key in selected_keys:
                    continue
                pick = cand
                break
        if pick is None:
            break
        assigned = _assign_backend(pick, preferred if preferred != "any" else "clawrag")
        selected_by_backend[assigned] += 1
        selected.append(pick)
        selected_keys.add((str(pick.get("path") or ""), int(pick.get("chunk_index") or 0)))

    selected_by_source_group: dict[str, int] = {}
    selected_by_origin: dict[str, int] = {}
    for hit in selected:
        sg = str(hit.get("source_group") or "general").strip() or "general"
        selected_by_source_group[sg] = selected_by_source_group.get(sg, 0) + 1
        origin = str(hit.get("origin") or "unknown").strip() or "unknown"
        selected_by_origin[origin] = selected_by_origin.get(origin, 0) + 1

    report = {
        "final_k": final_k,
        "rag_target": rag_target,
        "local_target": local_target,
        "selected_rag": selected_by_backend["clawrag"],
        "selected_local": selected_by_backend["local"],
        "candidates_total": len(candidates),
        "candidates_rag": len(rag_hits),
        "candidates_local": len(local_hits),
        "has_glossary": has_glossary,
        "glossary_needed": glossary_needed,
        "forced_glossary": forced_glossary,
        "selected_by_source_group": selected_by_source_group,
        "selected_by_origin": selected_by_origin,
    }

    # Final payload for model context: keep minimal fields.
    hits_out: list[dict[str, Any]] = []
    for hit in selected[:final_k]:
        hits_out.append(
            {
                "path": hit.get("path"),
                "source_group": hit.get("source_group") or "general",
                "chunk_index": int(hit.get("chunk_index") or 0),
                "snippet": str(hit.get("snippet") or "")[:700],
                "score": round(float(hit.get("final_score") or 0.0), 6),
                "origin": hit.get("origin") or "unknown",
            }
        )
    return hits_out, report


def _reference_like_filters(*, kb_root: Path, kb_company: str) -> tuple[str, str]:
    ref_root = (kb_root / "30_Reference").expanduser().resolve()
    company_root = (ref_root / kb_company).expanduser().resolve()
    return f"{ref_root}/%", f"{company_root}/%"


def _company_from_kb_path(path: str, *, kb_root: Path) -> str:
    """Infer company from KB path by structure: {Section}/{Company}/..."""
    p = Path(str(path)).expanduser().resolve()
    root = kb_root.expanduser().resolve()
    try:
        rel = p.relative_to(root)
    except ValueError:
        return ""
    parts = rel.parts
    if len(parts) < 2:
        return ""
    section = parts[0]
    if section not in {"00_Glossary", "10_Style_Guide", "20_Domain_Knowledge", "30_Reference", "40_Templates"}:
        return ""
    company = str(parts[1] or "").strip()
    if not company or company.startswith("."):
        return ""
    return company


def _company_like_filters(*, kb_root: Path, kb_company: str) -> list[str]:
    root = kb_root.expanduser().resolve()
    company = (kb_company or "").strip()
    if not company:
        return []
    roots = [
        root / "00_Glossary" / company,
        root / "10_Style_Guide" / company,
        root / "20_Domain_Knowledge" / company,
        root / "30_Reference" / company,
        root / "40_Templates" / company,
    ]
    return [f"{p}/%" for p in roots]


def _allow_kb_path(
    path: str,
    *,
    kb_root: Path | None,
    kb_company: str,
    isolation_mode: str,
) -> bool:
    p = str(path or "").strip()
    if not p:
        return False

    low = p.replace("\\", "/").lower()
    # Never allow raw source uploads under Reference in retrieval context.
    if "/30_reference/" in low and "/source/" in low:
        return False
    # Reference retrieval is "final-only" by design.
    if "/30_reference/" in low and "/final/" not in low:
        return False

    if kb_root:
        kb_root_abs = str(kb_root.expanduser().resolve())
        if not (p.startswith(kb_root_abs + "/") or p == kb_root_abs):
            return False
    try:
        if not Path(p).expanduser().exists():
            return False
    except OSError:
        return False

    if not kb_root or not kb_company.strip():
        return True

    mode = (isolation_mode or "company_strict").strip().lower()
    if mode not in {"reference_only", "all", "company_strict"}:
        mode = "company_strict"

    if mode == "company_strict":
        company = _company_from_kb_path(p, kb_root=kb_root)
        return company == kb_company.strip()

    ref_root = str((kb_root / "30_Reference").expanduser().resolve())
    company_root = str((kb_root / "30_Reference" / kb_company).expanduser().resolve())

    if mode == "all":
        return p.startswith(company_root + "/") or p == company_root

    # reference_only: only filter Reference; allow glossary/style/domain/templates globally
    if p.startswith(ref_root + "/"):
        return p.startswith(company_root + "/") or p == company_root
    return True


def _ensure_kb_fts(conn: sqlite3.Connection) -> bool:
    """Best-effort: enable FTS5 BM25 retrieval when supported by SQLite."""
    global _KB_FTS_AVAILABLE
    if _KB_FTS_AVAILABLE is False:
        return False

    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='kb_chunks_fts'"
    ).fetchone()
    if row:
        _KB_FTS_AVAILABLE = True
        return True

    try:
        conn.execute(
            "CREATE VIRTUAL TABLE kb_chunks_fts USING fts5(text, content='kb_chunks', content_rowid='id', tokenize='unicode61')"
        )
        conn.executescript(
            """
            CREATE TRIGGER IF NOT EXISTS kb_chunks_ai AFTER INSERT ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(rowid, text) VALUES (new.id, new.text);
            END;
            CREATE TRIGGER IF NOT EXISTS kb_chunks_ad AFTER DELETE ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(kb_chunks_fts, rowid, text) VALUES('delete', old.id, old.text);
            END;
            CREATE TRIGGER IF NOT EXISTS kb_chunks_au AFTER UPDATE ON kb_chunks BEGIN
              INSERT INTO kb_chunks_fts(kb_chunks_fts, rowid, text) VALUES('delete', old.id, old.text);
              INSERT INTO kb_chunks_fts(rowid, text) VALUES (new.id, new.text);
            END;
            """
        )
        conn.execute("INSERT INTO kb_chunks_fts(kb_chunks_fts) VALUES('rebuild')")
        conn.commit()
        _KB_FTS_AVAILABLE = True
        return True
    except sqlite3.OperationalError:
        _KB_FTS_AVAILABLE = False
        return False

try:  # Optional dependency
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional import
    PdfReader = None

try:  # Optional dependency
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional import
    load_workbook = None


def _normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 120) -> list[str]:
    norm = _normalize_text(text)
    if not norm:
        return []

    units = [u.strip() for u in re.split(r"(?:\n{2,}|(?<=[.!?])\s+)", norm) if u.strip()]
    chunks: list[str] = []
    cur = ""
    for unit in units:
        if not cur:
            cur = unit
            continue
        if len(cur) + 1 + len(unit) <= max_chars:
            cur = f"{cur} {unit}"
        else:
            chunks.append(cur)
            tail = cur[-overlap:] if overlap > 0 and len(cur) > overlap else cur
            cur = f"{tail} {unit}" if tail else unit
            if len(cur) > max_chars:
                chunks.append(cur[:max_chars])
                cur = cur[max_chars - overlap :] if overlap > 0 else ""
    if cur:
        chunks.append(cur)
    return [_normalize_text(c) for c in chunks if _normalize_text(c)]


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = _normalize_text(p.text)
        if t:
            lines.append(t)
    for t_idx, table in enumerate(doc.tables, start=1):
        lines.append(f"[Table {t_idx}]")
        for row in table.rows:
            cells = [_normalize_text(c.text).replace("\n", " / ") for c in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    pdftotext_bin = "/opt/homebrew/bin/pdftotext"
    if Path(pdftotext_bin).exists():
        try:
            proc = subprocess.run(
                [pdftotext_bin, "-layout", str(path), "-"],
                capture_output=True, text=True, timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return _normalize_text(proc.stdout)
        except (subprocess.TimeoutExpired, OSError):
            pass
    # Fallback to pypdf
    if PdfReader is None:
        raise RuntimeError("Neither pdftotext nor pypdf available for .pdf")
    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        txt = _normalize_text(page.extract_text() or "")
        if txt:
            lines.append(txt)
    return "\n".join(lines)


def _extract_text_plain(path: Path) -> str:
    return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))


def _extract_csv(path: Path) -> str:
    lines: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            cleaned = [_normalize_text(c) for c in row if _normalize_text(c)]
            if cleaned:
                lines.append(" | ".join(cleaned))
    return "\n".join(lines)


SHEETSMITH_SCRIPT = Path.home() / ".openclaw/workspace/skills/sheetsmith/scripts/sheetsmith.py"


def _extract_xlsx(path: Path) -> str:
    if SHEETSMITH_SCRIPT.exists():
        try:
            proc = subprocess.run(
                [sys.executable, str(SHEETSMITH_SCRIPT), "preview", str(path), "--rows", "9999"],
                capture_output=True, text=True, timeout=60,
            )
            if proc.returncode == 0 and proc.stdout.strip():
                return _normalize_text(proc.stdout)
        except (subprocess.TimeoutExpired, OSError):
            pass
    # Fallback to openpyxl
    if load_workbook is None:
        raise RuntimeError("Neither sheetsmith nor openpyxl available for .xlsx")
    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[Sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells: list[str] = []
            for value in row:
                if value is None:
                    continue
                text = _normalize_text(str(value))
                if text:
                    cells.append(text)
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def extract_text(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        return "docx", _extract_docx(path)
    if ext == ".pdf":
        return "pdf", _extract_pdf(path)
    if ext in {".md", ".txt"}:
        return ext[1:], _extract_text_plain(path)
    if ext == ".csv":
        return "csv", _extract_csv(path)
    if ext == ".xlsx":
        return "xlsx", _extract_xlsx(path)
    return "unknown", ""


def discover_kb_files(kb_root: Path) -> tuple[list[Path], list[str]]:
    files: list[Path] = []
    unscoped: list[str] = []
    for p in kb_root.rglob("*"):
        if not p.is_file():
            continue
        if p.name.startswith("~$") or p.name.startswith("."):
            continue
        low_path = str(p).lower().replace("\\", "/")
        # Do not index raw source uploads inside Reference projects by default.
        if "/30_reference/" in low_path and "/source/" in low_path:
            continue
        if "/30_reference/" in low_path and "/final/" not in low_path:
            continue
        if p.suffix.lower() not in KB_SUPPORTED_EXTENSIONS:
            continue
        company = _company_from_kb_path(str(p), kb_root=kb_root)
        if not company:
            try:
                rel = p.expanduser().resolve().relative_to(kb_root.expanduser().resolve())
                if rel.parts and rel.parts[0] in {"00_Glossary", "10_Style_Guide", "20_Domain_Knowledge", "30_Reference", "40_Templates"}:
                    unscoped.append(str(p.expanduser().resolve()))
            except Exception:
                pass
            continue
        files.append(p)
    files.sort(key=lambda x: str(x).lower())
    unscoped.sort(key=lambda x: x.lower())
    return files, unscoped


def sync_kb(
    *,
    conn: sqlite3.Connection,
    kb_root: Path,
    report_path: Path | None = None,
) -> dict[str, Any]:
    kb_root = kb_root.expanduser().resolve()
    files, unscoped = discover_kb_files(kb_root)

    existing_rows = conn.execute("SELECT * FROM kb_files").fetchall()
    existing = {row["path"]: dict(row) for row in existing_rows}
    seen_paths = set()

    report: dict[str, Any] = {
        "ok": True,
        "kb_root": str(kb_root),
        "scanned_count": len(files),
        "created": 0,
        "updated": 0,
        "metadata_only": 0,
        "metadata_only_paths": [],
        "unscoped_skipped": len(unscoped),
        "unscoped_skipped_paths": unscoped,
        "removed": 0,
        "removed_paths": [],
        "skipped": 0,
        "errors": [],
        "files": [],
        "indexed_at": utc_now_iso(),
    }

    for path in files:
        ap = str(path.resolve())
        seen_paths.add(ap)
        stat = path.stat()
        rec = existing.get(ap)
        try:
            if rec and int(rec["mtime_ns"]) == stat.st_mtime_ns and int(rec["size_bytes"]) == stat.st_size:
                report["skipped"] += 1
                continue

            sha = compute_sha256(path)
            if rec and rec.get("sha256") == sha:
                conn.execute(
                    """
                    UPDATE kb_files
                    SET mtime_ns=?, size_bytes=?, indexed_at=?
                    WHERE path=?
                    """,
                    (stat.st_mtime_ns, stat.st_size, utc_now_iso(), ap),
                )
                report["metadata_only"] += 1
                report["metadata_only_paths"].append(ap)
                continue

            parser, text = extract_text(path)
            chunks = _chunk_text(text)
            source_group = infer_source_group(path, kb_root)

            conn.execute("DELETE FROM kb_chunks WHERE path=?", (ap,))
            for idx, chunk in enumerate(chunks):
                conn.execute(
                    "INSERT INTO kb_chunks(path, source_group, chunk_index, text) VALUES(?,?,?,?)",
                    (ap, source_group, idx, chunk),
                )

            conn.execute(
                """
                INSERT INTO kb_files(path, mtime_ns, size_bytes, sha256, parser, source_group, chunk_count, indexed_at)
                VALUES(?,?,?,?,?,?,?,?)
                ON CONFLICT(path) DO UPDATE SET
                    mtime_ns=excluded.mtime_ns,
                    size_bytes=excluded.size_bytes,
                    sha256=excluded.sha256,
                    parser=excluded.parser,
                    source_group=excluded.source_group,
                    chunk_count=excluded.chunk_count,
                    indexed_at=excluded.indexed_at
                """,
                (ap, stat.st_mtime_ns, stat.st_size, sha, parser, source_group, len(chunks), utc_now_iso()),
            )

            report["files"].append(
                {
                    "path": ap,
                    "parser": parser,
                    "source_group": source_group,
                    "chunk_count": len(chunks),
                }
            )
            if rec:
                report["updated"] += 1
            else:
                report["created"] += 1
        except Exception as exc:  # pragma: no cover - keeps sync resilient
            report["errors"].append({"path": ap, "error": str(exc)})

    for old_path in existing:
        if old_path in seen_paths:
            continue
        conn.execute("DELETE FROM kb_chunks WHERE path=?", (old_path,))
        conn.execute("DELETE FROM kb_files WHERE path=?", (old_path,))
        report["removed"] += 1
        report["removed_paths"].append(str(old_path))

    conn.commit()
    report["ok"] = len(report["errors"]) == 0

    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json_dumps(report), encoding="utf-8")

    return report


def sync_kb_with_rag(
    *,
    conn: sqlite3.Connection,
    kb_root: Path,
    report_path: Path | None = None,
    rag_backend: str = "clawrag",
    rag_base_url: str = "http://127.0.0.1:8080",
    rag_collection: str = "translation-kb",
    rag_collection_mode: str = "auto",
    isolation_mode: str = "company_strict",
    focus_company: str = "",
) -> dict[str, Any]:
    local_report = sync_kb(conn=conn, kb_root=kb_root, report_path=report_path)
    rag_report: dict[str, Any] = {"ok": False, "backend": "local", "mode": "disabled", "sync": {}, "delete": {}}
    if str(rag_backend).strip().lower() == "clawrag":
        changed_paths = [str(x.get("path")) for x in (local_report.get("files") or []) if str(x.get("path", "")).strip()]
        changed_paths.extend([str(p) for p in (local_report.get("metadata_only_paths") or []) if str(p).strip()])
        changed_paths = sorted(set([str(p).strip() for p in changed_paths if str(p).strip()]))
        removed_paths = [str(p) for p in (local_report.get("removed_paths") or []) if str(p).strip()]
        mode_norm = (rag_collection_mode or "").strip().lower() or "auto"
        isolation_norm = (isolation_mode or "").strip().lower() or "company_strict"
        if mode_norm == "auto":
            mode_norm = "per_company" if isolation_norm == "company_strict" else "shared"

        if mode_norm in {"per_company", "per-company", "company", "company_scoped", "company-scoped"}:
            by_company_changed: dict[str, list[str]] = {}
            by_company_removed: dict[str, list[str]] = {}
            unscoped_changed: list[str] = []
            unscoped_removed: list[str] = []

            for p in changed_paths:
                company = _company_from_kb_path(p, kb_root=kb_root)
                if company:
                    by_company_changed.setdefault(company, []).append(p)
                else:
                    unscoped_changed.append(p)

            for p in removed_paths:
                company = _company_from_kb_path(p, kb_root=kb_root)
                if company:
                    by_company_removed.setdefault(company, []).append(p)
                else:
                    unscoped_removed.append(p)

            companies_set = set(by_company_changed.keys()) | set(by_company_removed.keys())
            focus = (focus_company or "").strip()
            if focus:
                companies_set.add(focus)
            companies = sorted(companies_set, key=lambda s: s.lower())
            per_company: list[dict[str, Any]] = []
            all_ok = True
            for company in companies:
                collection = resolve_rag_collection(
                    base_collection=rag_collection,
                    company=company,
                    mode="per_company",
                    isolation_mode=isolation_norm,
                )
                seeded = (
                    conn.execute(
                        "SELECT seeded_at FROM kb_rag_collections WHERE collection=?",
                        (collection,),
                    ).fetchone()
                    is not None
                )
                upload_paths = sorted(set(by_company_changed.get(company, [])))
                if not seeded:
                    likes = _company_like_filters(kb_root=kb_root, kb_company=company)
                    if likes:
                        where = " OR ".join(["path LIKE ?"] * len(likes))
                        rows = conn.execute(
                            f"SELECT path FROM kb_files WHERE {where}",
                            tuple(likes),
                        ).fetchall()
                        upload_paths = sorted(set([str(r["path"]) for r in rows if str(r["path"] or "").strip()]))
                delete_report = clawrag_delete(
                    removed_paths=sorted(set(by_company_removed.get(company, []))),
                    base_url=rag_base_url,
                    collection=collection,
                )
                sync_report = clawrag_sync(
                    changed_paths=upload_paths,
                    base_url=rag_base_url,
                    collection=collection,
                )
                ok = bool(sync_report.get("ok")) and bool(delete_report.get("ok"))
                if ok and not seeded:
                    conn.execute(
                        "INSERT OR REPLACE INTO kb_rag_collections(collection, mode, seeded_at) VALUES(?,?,?)",
                        (collection, "per_company", utc_now_iso()),
                    )
                all_ok = all_ok and ok
                per_company.append(
                    {
                        "company": company,
                        "collection": collection,
                        "seeded": seeded,
                        "uploaded_count": len(upload_paths),
                        "ok": ok,
                        "sync": sync_report,
                        "delete": delete_report,
                    }
                )

            # Best-effort cleanup for unscoped legacy paths (delete from base shared collection).
            legacy_delete = clawrag_delete(
                removed_paths=sorted(set(unscoped_removed)),
                base_url=rag_base_url,
                collection=(rag_collection or "translation-kb"),
            )
            legacy_sync = clawrag_sync(
                changed_paths=sorted(set(unscoped_changed)),
                base_url=rag_base_url,
                collection=(rag_collection or "translation-kb"),
            )
            if unscoped_changed or unscoped_removed:
                all_ok = all_ok and bool(legacy_delete.get("ok")) and bool(legacy_sync.get("ok"))

            conn.commit()
            rag_report = {
                "ok": bool(all_ok),
                "backend": "clawrag",
                "mode": "per_company",
                "base_collection": rag_collection,
                "companies": per_company,
                "legacy_unscoped": {
                    "changed_paths": unscoped_changed,
                    "removed_paths": unscoped_removed,
                    "sync": legacy_sync,
                    "delete": legacy_delete,
                },
            }
        else:
            delete_report = clawrag_delete(
                removed_paths=removed_paths,
                base_url=rag_base_url,
                collection=rag_collection,
            )
            sync_report = clawrag_sync(
                changed_paths=changed_paths,
                base_url=rag_base_url,
                collection=rag_collection,
            )
            rag_report = {
                "ok": bool(sync_report.get("ok")) and bool(delete_report.get("ok")),
                "backend": "clawrag",
                "mode": "shared",
                "collection": rag_collection,
                "sync": sync_report,
                "delete": delete_report,
            }
    return {"ok": local_report.get("ok", False), "local_report": local_report, "rag_report": rag_report}


def retrieve_kb(
    *,
    conn: sqlite3.Connection,
    query: str,
    task_type: str = "",
    top_k: int = 8,
    kb_root: Path | None = None,
    kb_company: str = "",
    isolation_mode: str = "company_strict",
) -> list[dict[str, Any]]:
    tokens = [t.lower() for t in re.findall(r"[A-Za-z0-9\u0600-\u06FF_]+", query) if len(t) >= 2]
    if not tokens:
        return []

    task = task_type.upper().strip()
    task_boosts = {
        "REVISION_UPDATE": {"glossary": 1.25, "previously_translated": 1.2},
        "NEW_TRANSLATION": {"glossary": 1.2, "arabic_source": 1.1},
        "BILINGUAL_REVIEW": {"glossary": 1.2, "translated_en": 1.15},
        "EN_ONLY_EDIT": {"translated_en": 1.2, "previously_translated": 1.1},
        "MULTI_FILE_BATCH": {"glossary": 1.15, "previously_translated": 1.1},
        "TERMINOLOGY_ENFORCEMENT": {"glossary": 1.4, "translated_en": 1.15},
        "FORMAT_CRITICAL_TASK": {"glossary": 1.2, "previously_translated": 1.15},
        "LOW_CONTEXT_TASK": {"glossary": 1.1},
    }.get(task, {})

    if _ensure_kb_fts(conn):
        match_query = " OR ".join(sorted(set(tokens)))
        where_sql = ""
        where_params: list[Any] = []
        if kb_root and kb_company.strip():
            mode = (isolation_mode or "company_strict").strip().lower()
            if mode == "company_strict":
                likes = _company_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
                where_sql = " AND (" + " OR ".join(["c.path LIKE ?"] * len(likes)) + ")"
                where_params.extend(likes)
            else:
                ref_like, company_like = _reference_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
                if mode == "all":
                    where_sql = " AND c.path LIKE ?"
                    where_params.append(company_like)
                else:
                    where_sql = " AND (c.path NOT LIKE ? OR c.path LIKE ?)"
                    where_params.extend([ref_like, company_like])

        rows = conn.execute(
            f"""
            SELECT
              c.path AS path,
              c.source_group AS source_group,
              c.chunk_index AS chunk_index,
              substr(c.text, 1, 700) AS snippet,
              bm25(kb_chunks_fts) AS rank
            FROM kb_chunks_fts
            JOIN kb_chunks c ON c.id = kb_chunks_fts.rowid
            WHERE kb_chunks_fts MATCH ?{where_sql}
            LIMIT ?
            """,
            [match_query, *where_params, max(10, int(top_k) * 8)],
        ).fetchall()

        scored: list[dict[str, Any]] = []
        for row in rows:
            source_group = row["source_group"] or "general"
            base = SOURCE_GROUP_WEIGHTS.get(source_group, SOURCE_GROUP_WEIGHTS["general"])
            boost = task_boosts.get(source_group, 1.0)
            raw_rank = float(row["rank"] or 0.0)
            inv = 1.0 / (1.0 + max(0.0, raw_rank))
            score = round(inv * float(base) * float(boost), 6)
            scored.append(
                {
                    "path": row["path"],
                    "source_group": source_group,
                    "chunk_index": int(row["chunk_index"]),
                    "snippet": str(row["snippet"] or ""),
                    "score": score,
                }
            )
        scored.sort(key=lambda x: x["score"], reverse=True)
        return scored[: max(1, int(top_k))]

    where_sql = ""
    where_params: list[Any] = []
    if kb_root and kb_company.strip():
        mode = (isolation_mode or "company_strict").strip().lower()
        if mode == "company_strict":
            likes = _company_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
            where_sql = " WHERE (" + " OR ".join(["path LIKE ?"] * len(likes)) + ")"
            where_params.extend(likes)
        else:
            ref_like, company_like = _reference_like_filters(kb_root=kb_root, kb_company=kb_company.strip())
            if mode == "all":
                where_sql = " WHERE path LIKE ?"
                where_params.append(company_like)
            else:
                where_sql = " WHERE (path NOT LIKE ? OR path LIKE ?)"
                where_params.extend([ref_like, company_like])

    rows = conn.execute(f"SELECT path, source_group, chunk_index, text FROM kb_chunks{where_sql}", tuple(where_params)).fetchall()
    scored: list[dict[str, Any]] = []
    for row in rows:
        text = (row["text"] or "").lower()
        if not text:
            continue
        match_hits = 0
        for tk in tokens:
            match_hits += text.count(tk)
        if match_hits <= 0:
            continue
        source_group = row["source_group"] or "general"
        base = SOURCE_GROUP_WEIGHTS.get(source_group, SOURCE_GROUP_WEIGHTS["general"])
        boost = task_boosts.get(source_group, 1.0)
        score = round(float(match_hits) * base * boost, 4)
        scored.append(
            {
                "path": row["path"],
                "source_group": source_group,
                "chunk_index": int(row["chunk_index"]),
                "snippet": row["text"][:700],
                "score": score,
            }
        )

    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[: max(1, int(top_k))]


def retrieve_kb_with_fallback(
    *,
    conn: sqlite3.Connection,
    query: str,
    task_type: str = "",
    kb_root: Path | None = None,
    kb_company: str = "",
    isolation_mode: str = "company_strict",
    rag_backend: str = "clawrag",
    rag_base_url: str = "http://127.0.0.1:8080",
    rag_collection: str = "translation-kb",
    top_k_clawrag: int = 20,
    top_k_local: int = 12,
) -> dict[str, Any]:
    q = (query or "").strip()
    if not q:
        return {"backend": "local", "hits": [], "status_flags": [], "rag_result": {"ok": True, "detail": "empty_query"}}

    if str(rag_backend).strip().lower() == "clawrag":
        status_flags: list[str] = []
        rag_result = clawrag_search(
            query=q,
            top_k=max(1, int(top_k_clawrag)),
            base_url=rag_base_url,
            collection=rag_collection,
        )
        rag_hits_raw = list(rag_result.get("hits") or [])
        rag_hits = [
            h
            for h in rag_hits_raw
            if _allow_kb_path(
                str(h.get("path") or ""),
                kb_root=kb_root,
                kb_company=kb_company,
                isolation_mode=isolation_mode,
            )
        ]
        if rag_result.get("ok") and rag_hits_raw and not rag_hits:
            status_flags.append("rag_filtered_empty")
        if not rag_result.get("ok"):
            status_flags.append("rag_fallback_local")

        local_hits = retrieve_kb(
            conn=conn,
            query=q,
            task_type=task_type,
            top_k=max(1, int(top_k_local)),
            kb_root=kb_root,
            kb_company=kb_company,
            isolation_mode=isolation_mode,
        )

        # Extra defense: apply the same allowlist rules to local hits.
        local_hits_filtered = [
            h
            for h in list(local_hits or [])
            if _allow_kb_path(
                str(h.get("path") or ""),
                kb_root=kb_root,
                kb_company=kb_company,
                isolation_mode=isolation_mode,
            )
        ]

        final_k = int(os.getenv("OPENCLAW_KB_RERANK_FINAL_K", "12"))
        glossary_min = int(os.getenv("OPENCLAW_KB_RERANK_GLOSSARY_MIN", "3"))
        terminology_ratio = float(os.getenv("OPENCLAW_KB_RERANK_TERMINOLOGY_GLOSSARY_RATIO", "0.4"))
        task_upper = (task_type or "").upper().strip()
        prefer_rag_ratio = 0.4 if task_upper == "TERMINOLOGY_ENFORCEMENT" else 0.6

        merged_hits, rerank_report = _merge_and_rerank_hits(
            rag_hits=rag_hits[: max(1, int(top_k_clawrag))],
            local_hits=local_hits_filtered[: max(1, int(top_k_local))],
            task_type=task_type,
            final_k=final_k,
            glossary_min=glossary_min,
            terminology_glossary_ratio=terminology_ratio,
            prefer_rag_ratio=prefer_rag_ratio,
        )

        backend = "merged" if rag_result.get("ok") else "local"
        status_flags = (["rag_merge_rerank"] if backend == "merged" else []) + status_flags
        return {
            "backend": backend,
            "hits": merged_hits,
            "status_flags": status_flags,
            "rag_result": {
                **dict(rag_result or {}),
                "top_k": int(top_k_clawrag),
                "rag_hits_raw_count": len(rag_hits_raw),
                "rag_hits_filtered_count": len(rag_hits),
                "local_hits_count": len(local_hits_filtered),
                "rerank_report": rerank_report,
            },
        }

    local_hits = retrieve_kb(
        conn=conn,
        query=q,
        task_type=task_type,
        top_k=max(1, int(top_k_local)),
        kb_root=kb_root,
        kb_company=kb_company,
        isolation_mode=isolation_mode,
    )
    return {"backend": "local", "hits": local_hits, "status_flags": [], "rag_result": {"ok": True, "mode": "local_only"}}
