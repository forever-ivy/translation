#!/usr/bin/env python3
"""Create a DOCX draft while preserving the template layout exactly."""

from __future__ import annotations

import argparse
import base64
import json
import shutil
import sys
from pathlib import Path

from docx import Document

def split_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def build_doc(template_path: Path, output_path: Path, text: str) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if template_path.resolve() == output_path.resolve():
        return
    shutil.copy2(template_path, output_path)
    if not split_lines(text):
        return

    doc = Document(str(output_path))
    replacements = split_lines(text)
    idx = 0

    def next_line() -> str | None:
        nonlocal idx
        if idx >= len(replacements):
            return None
        value = replacements[idx]
        idx += 1
        return value

    for p in doc.paragraphs:
        if idx >= len(replacements):
            break
        if not p.text.strip():
            continue
        candidate = next_line()
        if candidate is None:
            break
        p.text = candidate

    if idx < len(replacements):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if idx >= len(replacements):
                        break
                    if not cell.text.strip():
                        continue
                    candidate = next_line()
                    if candidate is None:
                        break
                    cell.text = candidate
                if idx >= len(replacements):
                    break
            if idx >= len(replacements):
                break

    while idx < len(replacements):
        doc.add_paragraph(replacements[idx])
        idx += 1

    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--draft-text", help="Plain text draft")
    parser.add_argument("--draft-text-file", help="Path to plain text draft")
    parser.add_argument("--draft-text-base64", help="Base64-encoded plain text draft")
    args = parser.parse_args()

    template_path = Path(args.template)
    output_path = Path(args.output)

    if not template_path.exists():
        print(json.dumps({"ok": False, "error": f"Missing template: {template_path}"}), file=sys.stderr)
        return 2

    text = ""
    if args.draft_text:
        text = args.draft_text
    elif args.draft_text_file:
        draft_file = Path(args.draft_text_file)
        if not draft_file.exists():
            print(json.dumps({"ok": False, "error": f"Missing draft text file: {draft_file}"}), file=sys.stderr)
            return 2
        text = draft_file.read_text(encoding="utf-8")
    elif args.draft_text_base64:
        text = base64.b64decode(args.draft_text_base64.encode("utf-8")).decode("utf-8")

    build_doc(template_path, output_path, text)

    print(
        json.dumps(
            {
                "ok": True,
                "data": {
                    "template": str(template_path.resolve()),
                    "output": str(output_path.resolve()),
                    "line_count": len(split_lines(text)),
                },
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
