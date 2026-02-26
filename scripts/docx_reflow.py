#!/usr/bin/env python3
"""DOCX reflow helpers.

Primary use: convert RTL/Arabic-templated documents into LTR English-friendly layout
after a format-preserving translation map has been applied.
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _iter_all_paragraphs(doc: Document) -> Iterable[Any]:
    # Top-level paragraphs (document body)
    for p in doc.paragraphs:
        yield p
    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _remove_all_children(parent: Any, child_tag: str) -> int:
    removed = 0
    try:
        tag = qn(child_tag)
    except Exception:
        tag = child_tag
    while True:
        try:
            child = parent.find(tag)
        except Exception:
            child = None
        if child is None:
            break
        try:
            parent.remove(child)
            removed += 1
        except Exception:
            break
    return removed


def reflow_docx_to_english(
    *,
    input_docx: Path,
    output_docx: Path,
    force_left_alignment: bool = True,
) -> dict[str, Any]:
    """Create a LTR-friendly copy of input_docx at output_docx."""
    input_docx = Path(input_docx).expanduser().resolve()
    output_docx = Path(output_docx).expanduser().resolve()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(input_docx), str(output_docx))

    doc = Document(str(output_docx))
    paragraph_align_fixed = 0
    paragraph_bidi_removed = 0
    run_rtl_fixed = 0

    for para in _iter_all_paragraphs(doc):
        try:
            ppr = para._p.get_or_add_pPr()  # type: ignore[attr-defined]
        except Exception:
            ppr = None
        if ppr is not None:
            paragraph_bidi_removed += _remove_all_children(ppr, "w:bidi")

        if force_left_alignment:
            try:
                alignment = para.alignment
            except Exception:
                alignment = None
            if alignment in (None, WD_ALIGN_PARAGRAPH.RIGHT):
                try:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph_align_fixed += 1
                except Exception:
                    pass

        for run in list(getattr(para, "runs", []) or []):
            try:
                if run.font.rtl is not False:  # type: ignore[attr-defined]
                    run.font.rtl = False  # type: ignore[attr-defined]
                    run_rtl_fixed += 1
            except Exception:
                continue

    doc.save(str(output_docx))
    return {
        "ok": True,
        "input": str(input_docx),
        "output": str(output_docx),
        "paragraph_align_fixed": int(paragraph_align_fixed),
        "paragraph_bidi_removed": int(paragraph_bidi_removed),
        "run_rtl_fixed": int(run_rtl_fixed),
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Path to input .docx")
    parser.add_argument("--output", required=True, help="Path to output .docx")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists() or input_path.suffix.lower() != ".docx":
        print(json.dumps({"ok": False, "error": f"Missing .docx input: {input_path}"}), file=sys.stderr)
        return 2

    result = reflow_docx_to_english(input_docx=input_path, output_docx=output_path)
    print(json.dumps(result, ensure_ascii=False))
    return 0 if result.get("ok") else 1


if __name__ == "__main__":
    raise SystemExit(main())

