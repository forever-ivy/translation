#!/usr/bin/env python3
"""Write V5.2 artifact bundle into the verify folder."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook

from scripts.compose_docx_from_draft import build_doc
from scripts.docx_preserver import (
    apply_translation_map as apply_docx_translation_map,
    extract_units as extract_docx_units,
)
from scripts.docx_reflow import reflow_docx_to_english
from scripts.xlsx_preserver import (
    apply_translation_map as apply_xlsx_translation_map,
    extract_translatable_cells,
)

SYSTEM_DIR_NAME = ".system"


def _write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(value, encoding="utf-8")


def _write_json(path: Path, value: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def _text_to_lines(text: str) -> list[str]:
    return [line.rstrip() for line in (text or "").splitlines()]


def _write_docx(path: Path, title: str, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue
        doc.add_paragraph(stripped)
    doc.save(str(path))


def _write_xlsx(path: Path, *, final_text: str, change_log_points: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Final"
    for idx, line in enumerate([ln for ln in final_text.splitlines() if ln.strip()], start=1):
        ws.cell(row=idx, column=1, value=line.strip())

    log = wb.create_sheet(title="ChangeLog")
    if change_log_points:
        for idx, row in enumerate(change_log_points, start=1):
            log.cell(row=idx, column=1, value=str(row))
    else:
        log.cell(row=1, column=1, value="No explicit change log points were returned by model.")
    wb.save(str(path))


def build_review_brief_lines(
    *,
    task_type: str,
    quality_report: dict[str, Any],
    status_flags: list[str],
    review_questions: list[str],
) -> list[str]:
    rounds = quality_report.get("rounds", [])
    convergence = bool(quality_report.get("convergence_reached"))
    stop_reason = quality_report.get("stop_reason", "unknown")
    lines = [
        f"Task type: {task_type}",
        f"Convergence reached: {convergence}",
        f"Stop reason: {stop_reason}",
        f"Status flags: {', '.join(status_flags) if status_flags else 'none'}",
        "",
        "## Round Results",
    ]
    if not rounds:
        lines.append("- No rounds produced.")
    else:
        for item in rounds:
            lines.append(
                f"- Round {item.get('round')}: pass={item.get('pass')} codex={item.get('codex_pass')} gemini={item.get('gemini_pass')}"
            )
            unresolved = item.get("unresolved") or []
            if unresolved:
                lines.append(f"-   unresolved: {', '.join(unresolved)}")

    lines.extend(["", "## Questions / Notes"])
    if review_questions:
        for q in review_questions:
            lines.append(f"- {q}")
    else:
        lines.append("- No extra notes.")
    lines.extend(
        [
            "",
            "## Manual Policy",
            "- Output is in _VERIFY only.",
            "- System will not auto-move files to final folder.",
            "- After manual validation, use command: ok (status only).",
        ]
    )
    return lines


def _ensure_change_log_text(change_log_points: list[str], task_type: str) -> str:
    if change_log_points:
        body = "\n".join(f"- {x}" for x in change_log_points if str(x).strip())
    else:
        body = "- No explicit change log points were returned by model."
    return "\n".join(
        [
            f"# Change Log ({task_type})",
            "",
            body,
            "",
            "## Delivery Note",
            "- This artifact is generated in _VERIFY only.",
            "- Manual move is required for final delivery.",
        ]
    )


def _normalize_docx_map_entries(entries: Any, *, default_file: str = "") -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    if not entries:
        return out
    if isinstance(entries, dict):
        if not default_file:
            return out
        for unit_id, text in entries.items():
            uid = str(unit_id or "").strip()
            if not uid:
                continue
            out.append({"file": default_file, "id": uid, "text": str(text or "")})
        return out
    if not isinstance(entries, list):
        return out
    for item in entries:
        if not isinstance(item, dict):
            continue
        unit_id = str(item.get("id") or item.get("unit_id") or item.get("block_id") or item.get("cell_id") or "").strip()
        file_name = str(item.get("file") or "").strip() or default_file
        if not unit_id:
            continue
        if not file_name and default_file:
            file_name = default_file
        out.append({"file": file_name, "id": unit_id, "text": str(item.get("text") or "")})
    return out


def _normalize_xlsx_map_entries(entries: Any, *, default_file: str = "") -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    if not isinstance(entries, list):
        return out
    for item in entries:
        if not isinstance(item, dict):
            continue
        file_name = str(item.get("file") or "").strip() or default_file
        sheet = str(item.get("sheet") or "").strip()
        cell = str(item.get("cell") or "").strip().upper()
        if not file_name and default_file:
            file_name = default_file
        if not sheet or not cell:
            continue
        out.append(
            {
                "file": file_name,
                "sheet": sheet,
                "cell": cell,
                "text": str(item.get("text") or ""),
            }
        )
    return out


def _build_docx_source_text_map(path: Path) -> dict[str, str]:
    try:
        units, _ = extract_docx_units(path, include_tables=True)
    except Exception:
        return {}
    return {str(unit.unit_id): str(unit.text or "") for unit in units}


def _build_xlsx_source_text_map(path: Path) -> dict[tuple[str, str], str]:
    try:
        units, _ = extract_translatable_cells(path)
    except Exception:
        return {}
    return {(str(unit.sheet), str(unit.cell).upper()): str(unit.text or "") for unit in units}


def _write_bilingual_docx(path: Path, *, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Bilingual", level=1)
    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "ID"
    header[1].text = "Source"
    header[2].text = "Target"
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get("id") or "")
        cells[1].text = str(row.get("source") or "")
        cells[2].text = str(row.get("target") or "")
    doc.save(str(path))


def _write_bilingual_xlsx(path: Path, *, rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Bilingual"
    ws.cell(row=1, column=1, value="Sheet")
    ws.cell(row=1, column=2, value="Cell")
    ws.cell(row=1, column=3, value="Source")
    ws.cell(row=1, column=4, value="Target")
    for idx, row in enumerate(rows, start=2):
        ws.cell(row=idx, column=1, value=str(row.get("sheet") or ""))
        ws.cell(row=idx, column=2, value=str(row.get("cell") or ""))
        ws.cell(row=idx, column=3, value=str(row.get("source") or ""))
        ws.cell(row=idx, column=4, value=str(row.get("target") or ""))
    wb.save(str(path))
    wb.close()


def _safe_unlink(path: Path) -> None:
    try:
        path.unlink(missing_ok=True)
    except Exception:
        return


def _cleanup_legacy_top_level_artifacts(review_dir: Path) -> None:
    for name in ("Final-Reflow.docx", "Review Brief.docx", "Change Log.md"):
        _safe_unlink(review_dir / name)


def write_artifacts(
    *,
    review_dir: str,
    draft_a_template_path: str | None,
    delta_pack: dict[str, Any],
    model_scores: dict[str, Any],
    quality: dict[str, Any],
    quality_report: dict[str, Any],
    job_id: str,
    task_type: str,
    confidence: float,
    estimated_minutes: int,
    runtime_timeout_minutes: int,
    iteration_count: int,
    double_pass: bool,
    status_flags: list[str],
    candidate_files: list[dict[str, Any]],
    review_questions: list[str],
    draft_payload: dict[str, Any] | None = None,
    generate_final_xlsx: bool = False,
    plan_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    review = Path(review_dir)
    system = review / SYSTEM_DIR_NAME
    review.mkdir(parents=True, exist_ok=True)
    system.mkdir(parents=True, exist_ok=True)
    _cleanup_legacy_top_level_artifacts(review)

    draft_payload = draft_payload or {}
    plan_payload = plan_payload or {}

    final_text = str(
        draft_payload.get("final_text")
        or draft_payload.get("draft_a_text")
        or draft_payload.get("draft_b_text")
        or ""
    )
    final_reflow_text = str(
        draft_payload.get("final_reflow_text")
        or draft_payload.get("draft_b_text")
        or final_text
        or ""
    )
    review_points = [str(x) for x in (draft_payload.get("review_brief_points") or [])]
    change_log_points = [str(x) for x in (draft_payload.get("change_log_points") or [])]
    docx_translation_map = draft_payload.get("docx_translation_map") or []
    xlsx_translation_map = draft_payload.get("xlsx_translation_map") or []

    final_docx = review / "Final.docx"
    bilingual_docx = review / "Bilingual.docx"
    final_xlsx = review / "Final.xlsx"
    bilingual_xlsx = review / "Bilingual.xlsx"
    final_reflow_docx = system / "final_reflow.docx"
    review_brief_md = system / "review_brief.md"
    change_log_md = system / "change_log.md"

    execution_plan_json = system / "execution_plan.json"
    quality_report_json = system / "quality_report.json"
    delta_summary_json = system / "Delta Summary.json"
    model_scores_json = system / "Model Scores.json"

    docx_sources = [
        Path(str(item.get("path") or "")).expanduser().resolve()
        for item in (candidate_files or [])
        if str(item.get("path") or "") and Path(str(item.get("path") or "")).suffix.lower() == ".docx"
    ]
    xlsx_sources = [
        Path(str(item.get("path") or "")).expanduser().resolve()
        for item in (candidate_files or [])
        if str(item.get("path") or "") and Path(str(item.get("path") or "")).suffix.lower() == ".xlsx"
    ]
    docx_default_file = docx_sources[0].name if len(docx_sources) == 1 else ""
    xlsx_default_file = xlsx_sources[0].name if len(xlsx_sources) == 1 else ""
    docx_map_rows = _normalize_docx_map_entries(docx_translation_map, default_file=docx_default_file)
    xlsx_map_rows = _normalize_xlsx_map_entries(xlsx_translation_map, default_file=xlsx_default_file)

    has_xlsx_outputs = bool(generate_final_xlsx or xlsx_sources)
    emit_docx = bool(docx_sources) or not has_xlsx_outputs

    template = Path(draft_a_template_path).expanduser() if draft_a_template_path else None

    docx_entries: list[dict[str, Any]] = []
    if emit_docx:
        final_docx_apply_result: dict[str, Any] = {}
        multi_docx_mode = len(docx_sources) >= 2 and bool(docx_map_rows)
        if not multi_docx_mode:
            if template and template.exists() and docx_map_rows:
                template_entries = [
                    {"id": row["id"], "text": row["text"]}
                    for row in docx_map_rows
                    if str(row.get("file") or "").strip() in {"", template.name}
                ]
                if not template_entries and len(docx_sources) == 1:
                    template_entries = [{"id": row["id"], "text": row["text"]} for row in docx_map_rows]
                final_docx_apply_result = apply_docx_translation_map(
                    template_docx=template,
                    output_docx=final_docx,
                    translation_map_entries=template_entries,
                )
            elif template and template.exists():
                build_doc(template, final_docx, final_text)
            else:
                _write_docx(final_docx, "Final", _text_to_lines(final_text))
        else:
            _safe_unlink(final_docx)
            _safe_unlink(bilingual_docx)

        if multi_docx_mode:
            for src in docx_sources:
                src_entries = [
                    {"id": row["id"], "text": row["text"]}
                    for row in docx_map_rows
                    if str(row.get("file") or "").strip() in {"", src.name}
                ]
                if not src_entries:
                    continue
                out_path = review / f"{src.stem}_translated.docx"
                bilingual_out = review / f"{src.stem}_bilingual.docx"
                res = apply_docx_translation_map(
                    template_docx=src,
                    output_docx=out_path,
                    translation_map_entries=src_entries,
                )
                source_map = _build_docx_source_text_map(src)
                rows = [
                    {"id": str(item["id"]), "source": str(source_map.get(str(item["id"]), "")), "target": str(item["text"])}
                    for item in src_entries
                ]
                if not rows:
                    rows = [{"id": "line:1", "source": "", "target": ""}]
                _write_bilingual_docx(bilingual_out, rows=rows)
                docx_entries.append(
                    {
                        "name": out_path.name,
                        "path": str(out_path.resolve()),
                        "source_path": str(src),
                        "apply_result": res,
                        "bilingual_path": str(bilingual_out.resolve()),
                    }
                )
        else:
            source_path = docx_sources[0] if docx_sources else template
            source_map = _build_docx_source_text_map(source_path) if source_path and source_path.exists() else {}
            if docx_map_rows:
                rows = [
                    {"id": str(item["id"]), "source": str(source_map.get(str(item["id"]), "")), "target": str(item["text"])}
                    for item in docx_map_rows
                ]
            else:
                lines = [ln.strip() for ln in final_text.splitlines() if ln.strip()]
                if not lines and source_map:
                    lines = [str(v) for _, v in sorted(source_map.items())]
                rows = [{"id": f"line:{idx}", "source": "", "target": line} for idx, line in enumerate(lines, start=1)]
            if not rows:
                rows = [{"id": "line:1", "source": "", "target": ""}]
            _write_bilingual_docx(bilingual_docx, rows=rows)
            docx_entries.append(
                {
                    "name": final_docx.name,
                    "path": str(final_docx.resolve()),
                    "source_path": str(source_path.resolve()) if source_path and source_path.exists() else "",
                    "apply_result": final_docx_apply_result,
                    "bilingual_path": str(bilingual_docx.resolve()),
                }
            )

        primary_docx = str(docx_entries[0]["path"]) if docx_entries else ""
        reflow_ok = False
        if primary_docx and Path(primary_docx).suffix.lower() == ".docx":
            try:
                reflow_res = reflow_docx_to_english(input_docx=Path(primary_docx), output_docx=final_reflow_docx)
                reflow_ok = bool(reflow_res.get("ok"))
            except Exception:
                reflow_ok = False
        if not reflow_ok:
            _write_docx(final_reflow_docx, "Final-Reflow", _text_to_lines(final_reflow_text))
    else:
        primary_docx = ""

    review_lines = build_review_brief_lines(
        task_type=task_type,
        quality_report=quality_report,
        status_flags=status_flags,
        review_questions=(review_points + review_questions),
    )
    _write_text(review_brief_md, "\n".join(review_lines))
    _write_text(change_log_md, _ensure_change_log_text(change_log_points, task_type))

    beautify_xlsx = str((plan_payload.get("meta") or {}).get("beautify_xlsx", "")).strip()
    if not beautify_xlsx:
        beautify_xlsx = "1"
    beautify_xlsx_enabled = beautify_xlsx not in {"0", "false", "off", "no"}

    xlsx_entries: list[dict[str, Any]] = []
    if generate_final_xlsx and xlsx_sources and xlsx_translation_map:
        if len(xlsx_sources) == 1:
            src = xlsx_sources[0]
            res = apply_xlsx_translation_map(
                source_xlsx=src,
                output_xlsx=final_xlsx,
                translation_map_entries=xlsx_translation_map,
                beautify=beautify_xlsx_enabled,
            )
            source_map = _build_xlsx_source_text_map(src)
            src_rows = [row for row in xlsx_map_rows if str(row.get("file") or "").strip() in {"", src.name}]
            rows = [
                {
                    "sheet": str(row.get("sheet") or ""),
                    "cell": str(row.get("cell") or ""),
                    "source": str(source_map.get((str(row.get("sheet") or ""), str(row.get("cell") or "").upper()), "")),
                    "target": str(row.get("text") or ""),
                }
                for row in src_rows
            ]
            if not rows:
                lines = [ln.strip() for ln in final_text.splitlines() if ln.strip()]
                rows = [{"sheet": "Final", "cell": f"A{idx}", "source": "", "target": line} for idx, line in enumerate(lines, start=1)]
            if not rows:
                rows = [{"sheet": "Final", "cell": "A1", "source": "", "target": ""}]
            _write_bilingual_xlsx(bilingual_xlsx, rows=rows)
            xlsx_entries.append(
                {
                    "name": final_xlsx.name,
                    "path": str(final_xlsx.resolve()),
                    "source_path": str(src),
                    "apply_result": res,
                    "bilingual_path": str(bilingual_xlsx.resolve()),
                }
            )
        else:
            for src in xlsx_sources:
                out_path = review / f"{src.stem}_translated.xlsx"
                bilingual_out = review / f"{src.stem}_bilingual.xlsx"
                res = apply_xlsx_translation_map(
                    source_xlsx=src,
                    output_xlsx=out_path,
                    translation_map_entries=xlsx_translation_map,
                    beautify=beautify_xlsx_enabled,
                )
                source_map = _build_xlsx_source_text_map(src)
                src_rows = [row for row in xlsx_map_rows if str(row.get("file") or "").strip() in {"", src.name}]
                rows = [
                    {
                        "sheet": str(row.get("sheet") or ""),
                        "cell": str(row.get("cell") or ""),
                        "source": str(source_map.get((str(row.get("sheet") or ""), str(row.get("cell") or "").upper()), "")),
                        "target": str(row.get("text") or ""),
                    }
                    for row in src_rows
                ]
                if not rows:
                    rows = [{"sheet": "Final", "cell": "A1", "source": "", "target": ""}]
                _write_bilingual_xlsx(bilingual_out, rows=rows)
                xlsx_entries.append(
                    {
                        "name": out_path.name,
                        "path": str(out_path.resolve()),
                        "source_path": str(src),
                        "apply_result": res,
                        "bilingual_path": str(bilingual_out.resolve()),
                    }
                )
    elif generate_final_xlsx:
        _write_xlsx(final_xlsx, final_text=final_text, change_log_points=change_log_points)
        lines = [ln.strip() for ln in final_text.splitlines() if ln.strip()]
        rows = [{"sheet": "Final", "cell": f"A{idx}", "source": "", "target": line} for idx, line in enumerate(lines, start=1)]
        if not rows:
            rows = [{"sheet": "Final", "cell": "A1", "source": "", "target": ""}]
        _write_bilingual_xlsx(bilingual_xlsx, rows=rows)
        xlsx_entries.append(
            {
                "name": final_xlsx.name,
                "path": str(final_xlsx.resolve()),
                "source_path": str(xlsx_sources[0]) if len(xlsx_sources) == 1 else "",
                "apply_result": {},
                "bilingual_path": str(bilingual_xlsx.resolve()),
            }
        )

    plan_write = {
        "job_id": job_id,
        "task_type": task_type,
        "confidence": confidence,
        "estimated_minutes": estimated_minutes,
        "runtime_timeout_minutes": runtime_timeout_minutes,
        "iteration_count": iteration_count,
        "double_pass": double_pass,
        "status_flags": status_flags,
        "candidate_files": candidate_files,
        "pipeline_version": str(
            (plan_payload.get("plan") or {}).get("pipeline_version")
            or (plan_payload.get("meta") or {}).get("pipeline_version")
            or ""
        ).strip(),
        "markdown_policy": (plan_payload.get("meta") or {}).get("markdown_policy") or {},
        "vision_policy": (plan_payload.get("meta") or {}).get("vision_policy") or {},
        "plan_payload": plan_payload,
    }
    _write_json(execution_plan_json, plan_write)
    _write_json(quality_report_json, quality_report)
    _write_json(delta_summary_json, delta_pack)
    _write_json(model_scores_json, model_scores)

    manifest: dict[str, Any] = {
        "primary_docx": str(primary_docx),
        "final_reflow_docx": str(final_reflow_docx.resolve()) if final_reflow_docx.exists() else "",
        "review_brief_md": str(review_brief_md.resolve()),
        "review_brief_docx": str(review_brief_md.resolve()),
        "change_log_md": str(change_log_md.resolve()),
        "execution_plan_json": str(execution_plan_json.resolve()),
        "quality_report_json": str(quality_report_json.resolve()),
        "delta_summary_json": str(delta_summary_json.resolve()),
        "model_scores_json": str(model_scores_json.resolve()),
    }
    if docx_entries:
        manifest["docx_files"] = docx_entries
        manifest["final_docx"] = str(docx_entries[0]["path"])
        bilingual_docx_path = str(docx_entries[0].get("bilingual_path") or "").strip()
        if bilingual_docx_path:
            manifest["bilingual_docx"] = bilingual_docx_path
    if xlsx_entries:
        manifest["xlsx_files"] = xlsx_entries
        if len(xlsx_entries) == 1 and xlsx_entries[0]["name"] == final_xlsx.name:
            manifest["final_xlsx"] = str(final_xlsx.resolve())
            bilingual_xlsx_path = str(xlsx_entries[0].get("bilingual_path") or "").strip()
            if bilingual_xlsx_path:
                manifest["bilingual_xlsx"] = bilingual_xlsx_path

    delivery_files: list[dict[str, Any]] = []
    for entry in docx_entries:
        path_value = str(entry.get("path") or "").strip()
        source_value = str(entry.get("source_path") or "").strip()
        if path_value:
            delivery_files.append(
                {"kind": "final_docx", "name": Path(path_value).name, "path": path_value, "source_path": source_value}
            )
        bilingual_value = str(entry.get("bilingual_path") or "").strip()
        if bilingual_value:
            delivery_files.append(
                {
                    "kind": "bilingual_docx",
                    "name": Path(bilingual_value).name,
                    "path": bilingual_value,
                    "source_path": source_value,
                }
            )
    for entry in xlsx_entries:
        path_value = str(entry.get("path") or "").strip()
        source_value = str(entry.get("source_path") or "").strip()
        if path_value:
            delivery_files.append(
                {"kind": "final_xlsx", "name": Path(path_value).name, "path": path_value, "source_path": source_value}
            )
        bilingual_value = str(entry.get("bilingual_path") or "").strip()
        if bilingual_value:
            delivery_files.append(
                {
                    "kind": "bilingual_xlsx",
                    "name": Path(bilingual_value).name,
                    "path": bilingual_value,
                    "source_path": source_value,
                }
            )
    seen_paths: set[str] = set()
    deduped: list[dict[str, Any]] = []
    for item in delivery_files:
        p = str(item.get("path") or "").strip()
        if not p or p in seen_paths:
            continue
        seen_paths.add(p)
        deduped.append(item)
    manifest["delivery_files"] = deduped
    return manifest
