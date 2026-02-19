#!/usr/bin/env python3
"""KB glossary parsing + strict terminology enforcement helpers.

This module is used by the translation orchestrator to enforce company-scoped
Arabic → English glossary translations.

Supported glossary sources (under KB root):
- 00_Glossary/<Company>/**/*.xlsx  (rows with Arabic + English columns)
- 00_Glossary/<Company>/**/*.docx  (tables or "Arabic - English" lines)

The orchestrator loads glossary pairs, then:
1) Selects only the pairs that appear in the job's source units/cells.
2) Hard-gates the run if any unit containing an Arabic glossary term does not
   contain the required English translation in the corresponding output.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

try:  # Optional dependency
    import openpyxl
except Exception:  # pragma: no cover
    openpyxl = None

try:  # Optional dependency
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

_ARABIC_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
_LATIN_RE = re.compile(r"[A-Za-z]")
_ARABIC_DIACRITICS_RE = re.compile(r"[\u064B-\u065F\u0670\u06D6-\u06ED]")
_WS_RE = re.compile(r"\s+")


@dataclass(frozen=True)
class GlossaryPair:
    arabic: str
    english: str
    source_path: str


def _normalize_space(text: str) -> str:
    return _WS_RE.sub(" ", (text or "").replace("\u00A0", " ")).strip()


def normalize_arabic(text: str) -> str:
    s = unicodedata.normalize("NFKC", text or "")
    s = s.replace("\u0640", "")  # tatweel
    s = _ARABIC_DIACRITICS_RE.sub("", s)
    return _normalize_space(s)


def normalize_english(text: str) -> str:
    s = unicodedata.normalize("NFKC", text or "")
    return _normalize_space(s).lower()


def _script_counts(text: str) -> tuple[int, int]:
    s = text or ""
    return (len(_ARABIC_RE.findall(s)), len(_LATIN_RE.findall(s)))


def looks_arabic(text: str) -> bool:
    ar, latin = _script_counts(text)
    return ar > 0 and ar >= latin


def looks_english(text: str) -> bool:
    ar, latin = _script_counts(text)
    return latin > 0 and latin >= ar


_HEADER_EN = {
    "english",
    "en",
    "translation",
    "english translation",
    "english term",
}
_HEADER_AR = {
    "العربية",
    "عربي",
    "العربي",
    "اللغة العربية",
}


def _is_header_pair(arabic: str, english: str) -> bool:
    return normalize_arabic(arabic) in _HEADER_AR and normalize_english(english) in _HEADER_EN


def _pick_best(candidates: list[str], *, prefer: str) -> str:
    if not candidates:
        return ""
    scored: list[tuple[int, int, str]] = []
    for c in candidates:
        ar, latin = _script_counts(c)
        scored.append((ar, latin, c))
    if prefer == "arabic":
        scored.sort(key=lambda x: (x[0], -x[1], len(x[2])), reverse=True)
    else:
        scored.sort(key=lambda x: (x[1], -x[0], len(x[2])), reverse=True)
    return scored[0][2]


def extract_pairs_from_xlsx(path: Path) -> list[GlossaryPair]:
    if openpyxl is None:  # pragma: no cover
        raise RuntimeError("openpyxl is required to parse .xlsx glossary files")

    path = Path(path).expanduser().resolve()
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    pairs: list[GlossaryPair] = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = []
                for v in row:
                    if v is None:
                        continue
                    if not isinstance(v, str):
                        v = str(v)
                    t = _normalize_space(v)
                    if t:
                        cells.append(t)
                if len(cells) < 2:
                    continue

                arabic_cells = [c for c in cells if looks_arabic(c)]
                english_cells = [c for c in cells if looks_english(c)]
                if not arabic_cells or not english_cells:
                    continue

                arabic = _pick_best(arabic_cells, prefer="arabic")
                english = _pick_best(english_cells, prefer="english")
                if not arabic or not english:
                    continue
                if _is_header_pair(arabic, english):
                    continue

                pairs.append(GlossaryPair(arabic=arabic, english=english, source_path=str(path)))
    finally:
        wb.close()
    return pairs


def _strip_bullet_prefix(text: str) -> str:
    s = (text or "").strip()
    s = re.sub(r"^[\s•\-\u2022\u25CF]+", "", s)
    s = re.sub(r"^\(?\d+[\)\.\-]\s*", "", s)  # "1) " / "1. "
    return s.strip()


def _split_glossary_line(line: str) -> tuple[str, str] | None:
    raw = _strip_bullet_prefix(line)
    if not raw:
        return None

    seps = [" | ", " - ", " – ", " — ", " —", ":", "=>", "→"]
    for sep in seps:
        if sep not in raw:
            continue
        left, right = raw.split(sep, 1)
        left = left.strip()
        right = right.strip()
        if not left or not right:
            continue
        if looks_arabic(left) and looks_english(right):
            return left, right
        if looks_english(left) and looks_arabic(right):
            return right, left
    return None


def extract_pairs_from_docx(path: Path) -> list[GlossaryPair]:
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is required to parse .docx glossary files")

    path = Path(path).expanduser().resolve()
    doc = Document(str(path))
    pairs: list[GlossaryPair] = []

    # Tables: preferred for bilingual glossary documents.
    for table in doc.tables:
        for row in table.rows:
            cells = [_normalize_space(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if len(cells) < 2:
                continue
            arabic_cells = [c for c in cells if looks_arabic(c)]
            english_cells = [c for c in cells if looks_english(c)]
            if not arabic_cells or not english_cells:
                continue
            arabic = _pick_best(arabic_cells, prefer="arabic")
            english = _pick_best(english_cells, prefer="english")
            if not arabic or not english:
                continue
            if _is_header_pair(arabic, english):
                continue
            pairs.append(GlossaryPair(arabic=arabic, english=english, source_path=str(path)))

    # Paragraph lines like "مدرسة - School" or "مدرسة: School"
    for para in doc.paragraphs:
        t = _normalize_space(para.text)
        if not t:
            continue
        maybe = _split_glossary_line(t)
        if not maybe:
            continue
        arabic, english = maybe
        if _is_header_pair(arabic, english):
            continue
        pairs.append(GlossaryPair(arabic=arabic, english=english, source_path=str(path)))

    return pairs


def load_company_glossary_pairs(
    *,
    kb_root: Path,
    company: str,
    max_files: int = 80,
) -> tuple[list[GlossaryPair], dict[str, Any]]:
    kb_root = Path(kb_root).expanduser().resolve()
    company_norm = (company or "").strip()
    glossary_dir = kb_root / "00_Glossary" / company_norm

    meta: dict[str, Any] = {
        "kb_root": str(kb_root),
        "company": company_norm,
        "glossary_dir": str(glossary_dir),
        "files_scanned": 0,
        "pairs_extracted": 0,
        "errors": [],
        "skipped_missing_dir": False,
    }

    if not glossary_dir.exists() or not glossary_dir.is_dir():
        meta["skipped_missing_dir"] = True
        return [], meta

    files: list[Path] = []
    for p in glossary_dir.rglob("*"):
        if not p.is_file():
            continue
        if p.name.startswith("~$") or p.name.startswith("."):
            continue
        if p.suffix.lower() not in {".xlsx", ".docx"}:
            continue
        files.append(p)

    files.sort(key=lambda p: str(p).lower())
    if max_files > 0:
        files = files[:max_files]

    pairs: list[GlossaryPair] = []
    for p in files:
        meta["files_scanned"] += 1
        try:
            if p.suffix.lower() == ".xlsx":
                pairs.extend(extract_pairs_from_xlsx(p))
            elif p.suffix.lower() == ".docx":
                pairs.extend(extract_pairs_from_docx(p))
        except Exception as exc:
            meta["errors"].append({"path": str(p), "error": str(exc)})

    meta["pairs_extracted"] = len(pairs)
    return pairs, meta


def build_glossary_map(
    pairs: Iterable[GlossaryPair],
    *,
    min_arabic_len: int = 2,
) -> tuple[dict[str, GlossaryPair], list[dict[str, Any]]]:
    """Return {arabic_norm: pair} plus conflict list."""
    out: dict[str, GlossaryPair] = {}
    conflicts: list[dict[str, Any]] = []

    for pair in pairs:
        ar = _normalize_space(pair.arabic)
        en = _normalize_space(pair.english)
        if not ar or not en:
            continue
        ar_norm = normalize_arabic(ar)
        if len(ar_norm) < max(1, int(min_arabic_len)):
            continue
        if not looks_arabic(ar_norm) or not looks_english(en):
            continue

        existing = out.get(ar_norm)
        if not existing:
            out[ar_norm] = GlossaryPair(arabic=ar, english=en, source_path=pair.source_path)
            continue

        if normalize_english(existing.english) == normalize_english(en):
            continue

        # Conflict: prefer longer English (more specific), otherwise keep first.
        chosen = existing
        if len(en) > len(existing.english):
            chosen = GlossaryPair(arabic=ar, english=en, source_path=pair.source_path)
            out[ar_norm] = chosen

        conflicts.append(
            {
                "arabic": ar,
                "arabic_norm": ar_norm,
                "existing_english": existing.english,
                "new_english": en,
                "kept_english": out[ar_norm].english,
                "existing_source": existing.source_path,
                "new_source": pair.source_path,
            }
        )

    return out, conflicts


def select_terms_for_sources(
    *,
    glossary_map: dict[str, GlossaryPair],
    source_texts: Iterable[str],
    max_terms: int = 80,
) -> tuple[list[GlossaryPair], dict[str, Any]]:
    """Select only glossary terms that appear in the provided source texts."""
    texts_norm = [normalize_arabic(_normalize_space(t)) for t in (source_texts or []) if _normalize_space(t)]
    meta: dict[str, Any] = {
        "source_text_count": len(texts_norm),
        "glossary_size": len(glossary_map),
        "max_terms": int(max_terms),
        "matched_terms": 0,
        "truncated": False,
    }
    if not texts_norm or not glossary_map:
        return [], meta

    term_hits: list[tuple[int, int, str]] = []
    for ar_norm, pair in glossary_map.items():
        hits = 0
        for t in texts_norm:
            if ar_norm and ar_norm in t:
                hits += 1
        if hits <= 0:
            continue
        term_hits.append((hits, len(ar_norm), ar_norm))

    term_hits.sort(key=lambda x: (x[0], x[1], x[2]), reverse=True)
    selected_norms = [ar_norm for _hits, _len, ar_norm in term_hits[: max(0, int(max_terms))]]

    if len(term_hits) > len(selected_norms):
        meta["truncated"] = True

    selected = [glossary_map[n] for n in selected_norms if n in glossary_map]
    meta["matched_terms"] = len(selected)
    return selected, meta

