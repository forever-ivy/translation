#!/usr/bin/env python3
"""PDF translation using PDFMathTranslate (pdf2zh) and Gemini Vision fallback."""

from __future__ import annotations

import base64
import json
import logging
import os
import shutil
import subprocess
import urllib.request
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)

# Environment variable names for different translation services
# pdf2zh-next uses these env vars
_SERVICE_ENV_VARS = {
    "google": None,  # Google Translate doesn't need API key (uses web scraping)
    "gemini": "GEMINI_API_KEY",
    "openai": "OPENAI_API_KEY",
    "deepl": "DEEPL_AUTH_KEY",
    "deepseek": "DEEPSEEK_API_KEY",
}

# Default locations to check for pdf2zh binary
_PDF2ZH_CANDIDATES = [
    shutil.which("pdf2zh"),  # In PATH
    Path.home() / ".local" / "bin" / "pdf2zh",  # uv/pip user install
    Path.home() / ".venv" / "pdf2zh" / "bin" / "pdf2zh",  # Isolated venv
]


def _find_pdf2zh() -> Path | None:
    """Find pdf2zh binary in common locations."""
    for candidate in _PDF2ZH_CANDIDATES:
        if candidate:
            p = Path(candidate)
            if p.exists() and p.is_file():
                return p
    return None


def is_pdf2zh_available() -> bool:
    """Check if pdf2zh is installed and available."""
    return _find_pdf2zh() is not None


def _find_ocrmypdf() -> Path | None:
    """Find ocrmypdf binary."""
    ocrmypdf_bin = shutil.which("ocrmypdf")
    if ocrmypdf_bin:
        return Path(ocrmypdf_bin)
    # Check common Homebrew path on macOS
    homebrew_path = Path("/opt/homebrew/bin/ocrmypdf")
    if homebrew_path.exists():
        return homebrew_path
    return None


def is_ocrmypdf_available() -> bool:
    """Check if ocrmypdf is installed and available."""
    return _find_ocrmypdf() is not None


def ocr_pdf(
    input_path: Path,
    output_path: Path,
    *,
    language: str = "ara",
    timeout_seconds: int = 300,
) -> dict[str, Any]:
    """
    Run OCR on a scanned PDF to add searchable text layer.

    Args:
        input_path: Path to source PDF file
        output_path: Path to write OCR'd PDF
        language: OCR language code (e.g., "ara" for Arabic, "eng" for English)
        timeout_seconds: Maximum time to wait for OCR

    Returns:
        dict with:
        - ok: bool - whether OCR succeeded
        - output_path: Path to OCR'd PDF
        - error: str - error message if failed
    """
    ocrmypdf_bin = _find_ocrmypdf()
    if not ocrmypdf_bin:
        return {"ok": False, "error": "ocrmypdf not installed. Install with: brew install ocrmypdf"}

    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        return {"ok": False, "error": f"Input file not found: {input_path}"}

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Force OCR to rasterize vector content and add text layer
    cmd = [
        str(ocrmypdf_bin),
        "--language", language,
        "--force-ocr",  # Force OCR even on vector content
        str(input_path),
        str(output_path),
    ]

    log.info("Running ocrmypdf: %s", " ".join(cmd))

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
        )
        if result.returncode != 0:
            error_msg = result.stderr.strip() or result.stdout.strip() or f"Exit code {result.returncode}"
            log.error("ocrmypdf failed: %s", error_msg)
            return {"ok": False, "error": f"ocrmypdf failed: {error_msg}"}
    except subprocess.TimeoutExpired:
        log.error("ocrmypdf timeout after %d seconds", timeout_seconds)
        return {"ok": False, "error": f"OCR timeout after {timeout_seconds}s"}
    except Exception as e:
        log.exception("ocrmypdf unexpected error")
        return {"ok": False, "error": str(e)}

    if not output_path.exists():
        return {"ok": False, "error": "OCR output file not created"}

    log.info("ocrmypdf success: %s", output_path)
    return {"ok": True, "output_path": output_path}


def translate_pdf(
    input_path: Path,
    output_dir: Path,
    *,
    service: str = "google",
    source_lang: str = "ar",
    target_lang: str = "en",
    timeout_seconds: int = 600,
    ocr_workaround: bool = True,
) -> dict[str, Any]:
    """
    Translate PDF using pdf2zh CLI.

    Args:
        input_path: Path to source PDF file
        output_dir: Directory to write translated PDFs
        service: Translation backend ("openai", "google", "deepl", etc.)
        source_lang: Source language code (e.g., "ar" for Arabic)
        target_lang: Target language code (e.g., "en" for English)
        timeout_seconds: Maximum time to wait for translation
        ocr_workaround: Enable OCR workaround for scanned PDFs

    Returns:
        dict with:
        - ok: bool - whether translation succeeded
        - mono_path: Path to translated PDF (target language only)
        - dual_path: Path to bilingual PDF (source + target side by side)
        - error: str - error message if failed
    """
    pdf2zh_bin = _find_pdf2zh()
    if not pdf2zh_bin:
        return {"ok": False, "error": "pdf2zh not installed. Install with: pip install pdf2zh"}

    input_path = Path(input_path)
    output_dir = Path(output_dir)

    if not input_path.exists():
        return {"ok": False, "error": f"Input file not found: {input_path}"}

    output_dir.mkdir(parents=True, exist_ok=True)

    # pdf2zh-next uses --google, --openai flags instead of -s service
    cmd = [
        str(pdf2zh_bin),
        str(input_path),
        "--output", str(output_dir),
        f"--{service}",  # e.g., --google, --openai
        "--lang-in", source_lang,
        "--lang-out", target_lang,
    ]

    # Add OCR workaround for scanned PDFs
    if ocr_workaround:
        cmd.append("--ocr-workaround")
        # Required for OCR'd/scanned PDFs to bypass scanned detection
        cmd.append("--skip-scanned-detection")

    log.info("Running pdf2zh: %s", " ".join(cmd))

    # Build environment with required API keys
    env = os.environ.copy()
    env_key = _SERVICE_ENV_VARS.get(service)
    if env_key and env_key in os.environ:
        # Already in environment, keep it
        pass
    elif env_key:
        # Try to load from .env.v4.local if not in environment
        env_path = Path(__file__).parent.parent / ".env.v4.local"
        if env_path.exists():
            for line in env_path.read_text().splitlines():
                if line.startswith(f"{env_key}="):
                    # Parse the value (handle quotes)
                    value = line.split("=", 1)[1].strip().strip('"').strip("'")
                    env[env_key] = value
                    log.info("Loaded %s from .env.v4.local", env_key)
                    break

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            env=env,
        )
        if result.returncode != 0:
            error_msg = result.stderr.strip() or result.stdout.strip() or f"Exit code {result.returncode}"
            log.error("pdf2zh failed: %s", error_msg)
            return {
                "ok": False,
                "error": f"pdf2zh failed: {error_msg}",
            }
    except subprocess.TimeoutExpired:
        log.error("pdf2zh timeout after %d seconds", timeout_seconds)
        return {"ok": False, "error": f"Translation timeout after {timeout_seconds}s"}
    except Exception as e:
        log.exception("pdf2zh unexpected error")
        return {"ok": False, "error": str(e)}

    # pdf2zh-next generates: {stem}.{target_lang}.mono.pdf and {stem}.{target_lang}.dual.pdf
    # Legacy pdf2zh generates: {stem}-mono.pdf and {stem}-dual.pdf
    stem = input_path.stem
    # Try new naming pattern first (pdf2zh-next)
    mono_path = output_dir / f"{stem}.{target_lang}.mono.pdf"
    dual_path = output_dir / f"{stem}.{target_lang}.dual.pdf"
    # Also try legacy pattern
    mono_path_legacy = output_dir / f"{stem}-mono.pdf"
    dual_path_legacy = output_dir / f"{stem}-dual.pdf"

    # Verify outputs exist
    outputs = {}
    for mono_candidate in [mono_path, mono_path_legacy]:
        if mono_candidate.exists():
            outputs["mono_path"] = mono_candidate
            break
    for dual_candidate in [dual_path, dual_path_legacy]:
        if dual_candidate.exists():
            outputs["dual_path"] = dual_candidate
            break

    if not outputs:
        # Check for alternative naming (some versions use different patterns)
        alt_files = list(output_dir.glob(f"{stem}*.pdf"))
        if alt_files:
            for f in alt_files:
                if "mono" in f.name.lower():
                    outputs["mono_path"] = f
                elif "dual" in f.name.lower() or "bilingual" in f.name.lower():
                    outputs["dual_path"] = f
                else:
                    # Default unknown files to mono
                    outputs.setdefault("mono_path", f)

    if not outputs:
        return {"ok": False, "error": "No output PDF generated"}

    log.info("pdf2zh success: %s", outputs)
    return {"ok": True, **outputs}


def translate_pdf_fallback_text(
    input_path: Path,
    output_dir: Path,
    *,
    source_lang: str = "ar",
    extract_func=None,
) -> dict[str, Any]:
    """
    Fallback: extract text from PDF without layout preservation.

    Uses the existing v4_kb._extract_pdf function for text extraction.
    If direct extraction returns empty text and ocrmypdf is available,
    attempts OCR before giving up.
    Output is written as a .txt file.

    Args:
        input_path: Path to source PDF file
        output_dir: Directory to write extracted text
        source_lang: Source language code for OCR (e.g., "ar" for Arabic)
        extract_func: Function to extract text (injected to avoid circular import)

    Returns:
        dict with:
        - ok: bool - whether extraction succeeded
        - text_path: Path to extracted text file
        - ocr_used: bool - whether OCR was used
        - error: str - error message if failed
    """
    from scripts.v4_kb import _extract_pdf as default_extract

    extract_fn = extract_func or default_extract
    input_path = Path(input_path)
    output_dir = Path(output_dir)

    if not input_path.exists():
        return {"ok": False, "error": f"Input file not found: {input_path}"}

    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        text = extract_fn(input_path)
        if text and text.strip():
            text_path = output_dir / f"{input_path.stem}-extracted.txt"
            text_path.write_text(text, encoding="utf-8")
            return {"ok": True, "text_path": text_path, "text": text, "ocr_used": False}

        # Direct extraction returned empty — try OCR if available
        if not is_ocrmypdf_available():
            return {"ok": False, "error": "No text extracted from PDF (ocrmypdf not installed for OCR fallback)"}

        log.info("Direct text extraction empty, attempting OCR for %s", input_path.name)
        tesseract_lang = _LANG_TO_TESSERACT.get(source_lang, source_lang)
        ocr_output_path = output_dir / f"{input_path.stem}_ocr_temp.pdf"

        ocr_result = ocr_pdf(input_path, ocr_output_path, language=tesseract_lang)
        if not ocr_result.get("ok"):
            return {"ok": False, "error": f"No text extracted and OCR failed: {ocr_result.get('error')}"}

        text = extract_fn(ocr_output_path)
        # Clean up temp OCR file
        try:
            ocr_output_path.unlink(missing_ok=True)
        except Exception:
            pass

        if not text or not text.strip():
            return {"ok": False, "error": "No text extracted from PDF even after OCR"}

        text_path = output_dir / f"{input_path.stem}-extracted.txt"
        text_path.write_text(text, encoding="utf-8")
        return {"ok": True, "text_path": text_path, "text": text, "ocr_used": True}
    except Exception as e:
        log.exception("PDF text extraction failed")
        return {"ok": False, "error": str(e)}


def check_pdf2zh_installation() -> dict[str, Any]:
    """
    Check pdf2zh installation status and provide installation hints.

    Returns:
        dict with installation status and helpful messages
    """
    pdf2zh_bin = _find_pdf2zh()
    if pdf2zh_bin:
        return {
            "installed": True,
            "path": str(pdf2zh_bin),
            "message": "pdf2zh is installed and ready",
        }

    return {
        "installed": False,
        "path": None,
        "message": "pdf2zh not found. Install with one of:\n"
        "  pip install pdf2zh\n"
        "  uv tool install --python 3.12 pdf2zh",
    }


# Error patterns that indicate a scanned PDF needs OCR
_SCANNED_PDF_ERROR_PATTERNS = [
    "no paragraphs",
    "Scanned PDF detected",
    "no text",
    "contains no extractable text",
    "No output PDF generated",
]


def _is_scanned_pdf_error(error_msg: str) -> bool:
    """Check if an error message indicates a scanned PDF issue."""
    error_lower = error_msg.lower()
    return any(pattern.lower() in error_lower for pattern in _SCANNED_PDF_ERROR_PATTERNS)


# Language code mapping: pdf2zh -> tesseract/ocrmypdf
_LANG_TO_TESSERACT = {
    "ar": "ara",
    "en": "eng",
    "zh": "chi_sim",
    "ja": "jpn",
    "ko": "kor",
    "fr": "fra",
    "de": "deu",
    "es": "spa",
    "ru": "rus",
    "pt": "por",
    "it": "ita",
}


# Human-readable language names for vision prompts
_LANG_NAMES = {
    "ar": "Arabic", "en": "English", "zh": "Chinese", "ja": "Japanese",
    "ko": "Korean", "fr": "French", "de": "German", "es": "Spanish",
    "ru": "Russian", "pt": "Portuguese", "it": "Italian",
}


def _pdf_to_pngs(pdf_path: Path, output_dir: Path, *, dpi: int = 200) -> list[Path]:
    """Convert each PDF page to a PNG using pdftoppm."""
    output_dir.mkdir(parents=True, exist_ok=True)
    prefix = output_dir / "page"
    cmd = ["pdftoppm", "-png", "-r", str(dpi), str(pdf_path), str(prefix)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=120)
    except FileNotFoundError as exc:
        raise RuntimeError(
            "pdftoppm not found. Install poppler: brew install poppler"
        ) from exc
    pages = sorted(output_dir.glob("page-*.png"), key=lambda p: p.name)
    if not pages:
        raise RuntimeError(f"pdftoppm produced no PNGs for {pdf_path.name}")
    return pages


def _load_env_key(key_name: str) -> str:
    """Load an env var, falling back to .env.v4.local if not in environment."""
    value = os.environ.get(key_name, "").strip()
    if value:
        return value
    env_path = Path(__file__).parent.parent / ".env.v4.local"
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            if line.startswith(f"{key_name}="):
                return line.split("=", 1)[1].strip().strip('"').strip("'")
    return ""


def _call_gemini_vision(image_b64: str, prompt_text: str, *, timeout: int = 120) -> str:
    """Send a single image + prompt to Gemini Vision and return the text response."""
    api_key = _load_env_key("GOOGLE_API_KEY") or _load_env_key("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY (or GEMINI_API_KEY) required for vision translation")
    model = (_load_env_key("OPENCLAW_GEMINI_VISION_MODEL") or "gemini-2.5-flash").strip()
    if "/" in model:
        model = model.rsplit("/", 1)[-1]
    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/{model}"
        f":generateContent?key={api_key}"
    )
    payload = {
        "contents": [{
            "parts": [
                {"text": prompt_text},
                {"inline_data": {"mime_type": "image/png", "data": image_b64}},
            ],
        }],
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    text = (
        ((body.get("candidates") or [{}])[0].get("content") or {})
        .get("parts", [{}])[0]
        .get("text", "")
    )
    return text


def _extract_first_json_object(raw: str) -> dict[str, Any]:
    """Extract the first JSON object from mixed model text output."""
    decoder = json.JSONDecoder()
    for idx, ch in enumerate(raw):
        if ch != "{":
            continue
        try:
            value, _ = decoder.raw_decode(raw[idx:])
            if isinstance(value, dict):
                return value
        except json.JSONDecodeError:
            continue
    raise ValueError("no JSON object found in response text")


def translate_pdf_vision(
    input_path: Path,
    output_dir: Path,
    *,
    source_lang: str = "ar",
    target_lang: str = "en",
    dpi: int = 200,
    timeout_per_page: int = 120,
) -> dict[str, Any]:
    """
    Translate a PDF using Gemini Vision (page-by-page image translation).

    Each page is rendered to PNG, sent to Gemini Vision which reads the text
    directly from the image and returns a structured translation.

    Returns:
        dict with:
        - ok: bool
        - docx_path: Path to the translated DOCX
        - page_count: int
        - error: str (if failed)
    """
    input_path = Path(input_path)
    output_dir = Path(output_dir)

    if not input_path.exists():
        return {"ok": False, "error": f"Input file not found: {input_path}"}

    api_key = _load_env_key("GOOGLE_API_KEY") or _load_env_key("GEMINI_API_KEY")
    if not api_key:
        return {"ok": False, "error": "GOOGLE_API_KEY not set — cannot use vision translation"}

    output_dir.mkdir(parents=True, exist_ok=True)
    png_dir = output_dir / "_vision_pages"

    try:
        pages = _pdf_to_pngs(input_path, png_dir, dpi=dpi)
    except RuntimeError as e:
        return {"ok": False, "error": str(e)}

    src_name = _LANG_NAMES.get(source_lang, source_lang)
    tgt_name = _LANG_NAMES.get(target_lang, target_lang)

    prompt_text = (
        f"This is a scanned document page in {src_name}. "
        f"Translate ALL text on this page into {tgt_name}.\n\n"
        "Return ONLY a JSON object with:\n"
        '- "translated_text": the full English translation of all text on the page\n'
        '- "sections": a list of objects, each with:\n'
        '  - "type": one of "header", "field", "paragraph", "list", "signature", "table"\n'
        '  - "content": the translated text for that section\n\n'
        "Rules:\n"
        "- Preserve the document structure (headers, numbered clauses, fields, signature blocks)\n"
        "- For form fields, use the format \"Label: Value\"\n"
        "- For numbered lists/clauses, keep the numbering\n"
        "- For tables, represent as rows of \"Column | Column\" separated by pipes\n"
        "- Translate everything — do not leave any text in the original language\n"
    )

    page_translations: list[dict[str, Any]] = []
    failed_pages = 0
    last_error = ""
    for i, png_path in enumerate(pages, start=1):
        log.info("Vision translating page %d/%d of %s", i, len(pages), input_path.name)
        img_b64 = base64.b64encode(png_path.read_bytes()).decode("ascii")
        try:
            raw_response = _call_gemini_vision(img_b64, prompt_text, timeout=timeout_per_page)
            page_data = _extract_first_json_object(raw_response)
        except Exception as e:
            log.warning("Vision translation failed for page %d: %s", i, e)
            failed_pages += 1
            last_error = str(e)
            page_data = {
                "translated_text": f"[Page {i} translation failed: {e}]",
                "sections": [{"type": "paragraph", "content": f"[Page {i} translation failed: {e}]"}],
            }
        page_translations.append(page_data)

    # If ALL pages failed, don't produce a bogus DOCX
    if failed_pages == len(pages):
        # Clean up PNGs
        try:
            shutil.rmtree(png_dir, ignore_errors=True)
        except Exception:
            pass
        return {"ok": False, "error": f"Vision translation failed for all {len(pages)} pages: {last_error}"}

    # Build DOCX from structured translations
    try:
        docx_path = _build_translated_docx(page_translations, output_dir, stem=input_path.stem)
    except Exception as e:
        log.exception("Failed to build DOCX from vision translations")
        return {"ok": False, "error": f"DOCX build failed: {e}"}

    # Clean up PNGs
    try:
        shutil.rmtree(png_dir, ignore_errors=True)
    except Exception:
        pass

    return {"ok": True, "docx_path": docx_path, "page_count": len(pages), "vision_used": True}


def _build_translated_docx(
    page_translations: list[dict[str, Any]],
    output_dir: Path,
    *,
    stem: str = "translated",
) -> Path:
    """Build a styled DOCX from per-page structured translations."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for page_idx, page_data in enumerate(page_translations):
        if page_idx > 0:
            doc.add_page_break()

        sections = page_data.get("sections") or []
        if not sections:
            # Fallback: use translated_text as a single paragraph
            text = page_data.get("translated_text", "")
            if text:
                doc.add_paragraph(text)
            continue

        for section in sections:
            sec_type = str(section.get("type", "paragraph")).lower()
            content = str(section.get("content", ""))
            if not content:
                continue

            if sec_type == "header":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(content)
                run.bold = True
                run.font.size = Pt(14)

            elif sec_type == "field":
                p = doc.add_paragraph()
                # Fields are "Label: Value" pairs
                if ":" in content:
                    label, _, value = content.partition(":")
                    run_label = p.add_run(label.strip() + ": ")
                    run_label.bold = True
                    p.add_run(value.strip())
                else:
                    p.add_run(content)

            elif sec_type == "list":
                # Handle multi-line numbered/bulleted lists
                for line in content.split("\n"):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line, style="List Number")

            elif sec_type == "table":
                # Parse pipe-separated table rows
                rows = [r.strip() for r in content.split("\n") if r.strip()]
                if rows:
                    cols = max(len(r.split("|")) for r in rows)
                    cols = max(cols, 1)
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Table Grid"
                    for r_idx, row_text in enumerate(rows):
                        cells = [c.strip() for c in row_text.split("|")]
                        for c_idx, cell_text in enumerate(cells[:cols]):
                            table.rows[r_idx].cells[c_idx].text = cell_text

            elif sec_type == "signature":
                doc.add_paragraph("")  # spacing
                p = doc.add_paragraph()
                p.add_run(content)
                doc.add_paragraph("_" * 40)
                doc.add_paragraph("")  # spacing

            else:
                # paragraph or unknown type
                doc.add_paragraph(content)

    docx_path = output_dir / f"{stem}-Vision-Translated.docx"
    doc.save(str(docx_path))
    return docx_path


def translate_pdf_with_ocr_fallback(
    input_path: Path,
    output_dir: Path,
    *,
    service: str = "google",
    source_lang: str = "ar",
    target_lang: str = "en",
    timeout_seconds: int = 600,
    ocr_timeout_seconds: int = 300,
) -> dict[str, Any]:
    """
    Translate PDF with automatic OCR fallback for scanned documents.

    This is the recommended entry point for PDF translation. It:
    1. First tries pdf2zh directly
    2. If pdf2zh fails due to scanned PDF issues, prefers Gemini Vision
       (renders pages as images → readable translated DOCX)
    3. Falls back to pdf2zh on OCR'd PDF only if Vision is unavailable

    Args:
        input_path: Path to source PDF file
        output_dir: Directory to write translated PDFs
        service: Translation backend ("google", "openai", "deepl", etc.)
        source_lang: Source language code (e.g., "ar" for Arabic)
        target_lang: Target language code (e.g., "en" for English)
        timeout_seconds: Maximum time to wait for translation
        ocr_timeout_seconds: Maximum time to wait for OCR preprocessing

    Returns:
        dict with:
        - ok: bool - whether translation succeeded
        - mono_path: Path to translated PDF (target language only)
        - dual_path: Path to bilingual PDF (source + target side by side)
        - ocr_used: bool - whether OCR preprocessing was used
        - error: str - error message if failed
    """
    input_path = Path(input_path)
    output_dir = Path(output_dir)

    # First attempt: try pdf2zh directly
    result = translate_pdf(
        input_path,
        output_dir,
        service=service,
        source_lang=source_lang,
        target_lang=target_lang,
        timeout_seconds=timeout_seconds,
        ocr_workaround=True,
    )

    if result.get("ok"):
        result["ocr_used"] = False
        return result

    error_msg = result.get("error", "")

    # Check if this is a scanned PDF issue that OCR might fix
    if not _is_scanned_pdf_error(error_msg):
        # Not a scanned PDF issue, return the original error
        return result

    # Check if ocrmypdf is available
    if not is_ocrmypdf_available():
        log.warning("Scanned PDF detected but ocrmypdf not available for OCR preprocessing")
        return {
            **result,
            "error": f"{error_msg}\n(ocrmypdf not installed - required for scanned PDFs)",
        }

    log.info("Scanned PDF detected, running OCR preprocessing with ocrmypdf")

    # Get tesseract language code
    tesseract_lang = _LANG_TO_TESSERACT.get(source_lang, source_lang)

    # Create temp path for OCR'd PDF
    ocr_output_path = output_dir / f"{input_path.stem}_ocr_temp.pdf"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Run OCR
    ocr_result = ocr_pdf(
        input_path,
        ocr_output_path,
        language=tesseract_lang,
        timeout_seconds=ocr_timeout_seconds,
    )

    if not ocr_result.get("ok"):
        log.error("OCR preprocessing failed: %s", ocr_result.get("error"))
        return {
            **result,
            "error": f"Scanned PDF detected and OCR preprocessing failed: {ocr_result.get('error')}",
        }

    # Scanned PDF: pdf2zh only translates the text layer, not the images.
    # The visual content stays in the original language. Prefer Gemini Vision
    # which renders pages as images and produces a readable translated DOCX.
    vision_api_key = _load_env_key("GOOGLE_API_KEY") or _load_env_key("GEMINI_API_KEY")
    if vision_api_key:
        log.info("Scanned PDF detected, using Gemini Vision translation for %s", input_path.name)
        # Clean up temp OCR file before vision path
        try:
            ocr_output_path.unlink(missing_ok=True)
        except Exception:
            pass
        vision_result = translate_pdf_vision(
            input_path,
            output_dir,
            source_lang=source_lang,
            target_lang=target_lang,
        )
        if vision_result.get("ok"):
            vision_result["ocr_used"] = False
            return vision_result
        log.warning("Vision translation failed: %s", vision_result.get("error"))

    # Vision not available or failed — fall back to pdf2zh on OCR'd PDF
    log.info("Falling back to pdf2zh on OCR'd PDF for %s", input_path.name)
    retry_result = translate_pdf(
        ocr_output_path,
        output_dir,
        service=service,
        source_lang=source_lang,
        target_lang=target_lang,
        timeout_seconds=timeout_seconds,
        ocr_workaround=True,
    )

    # Clean up temp OCR file
    try:
        ocr_output_path.unlink(missing_ok=True)
    except Exception:
        pass

    if retry_result.get("ok"):
        retry_result["ocr_used"] = True
        return retry_result

    # All attempts failed
    return {
        "ok": False,
        "error": f"Translation failed even after OCR preprocessing: {retry_result.get('error')}",
        "ocr_used": True,
    }
