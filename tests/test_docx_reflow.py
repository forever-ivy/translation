#!/usr/bin/env python3

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement

from scripts.docx_reflow import reflow_docx_to_english


class DocxReflowTest(unittest.TestCase):
    def test_reflow_removes_bidi_and_forces_ltr_runs(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "src.docx"
            out = Path(tmp) / "out.docx"

            doc = Document()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            ppr = p._p.get_or_add_pPr()  # type: ignore[attr-defined]
            ppr.append(OxmlElement("w:bidi"))

            run = p.add_run("Hello")
            run.font.rtl = True
            doc.save(src)

            res = reflow_docx_to_english(input_docx=src, output_docx=out)
            self.assertTrue(res.get("ok"))
            self.assertTrue(out.exists())

            doc2 = Document(str(out))
            p2 = doc2.paragraphs[0]
            self.assertEqual(p2.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            self.assertNotIn("<w:bidi", p2._p.xml)  # type: ignore[attr-defined]
            self.assertEqual(p2.runs[0].font.rtl, False)


if __name__ == "__main__":
    raise SystemExit(unittest.main())

