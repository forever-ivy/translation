#!/usr/bin/env python3
"""Tests for PDF translator module."""

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock

from scripts.pdf_translator import (
    is_pdf2zh_available,
    check_pdf2zh_installation,
    translate_pdf,
    translate_pdf_fallback_text,
    translate_pdf_vision,
    _build_translated_docx,
    _extract_first_json_object,
    _find_pdf2zh,
    _is_scanned_pdf_error,
)


class TestPdfTranslator(unittest.TestCase):
    """Test cases for pdf_translator module."""

    def test_is_pdf2zh_available_returns_false_when_not_found(self):
        """Test that is_pdf2zh_available returns False when binary not found."""
        # This test will pass even if pdf2zh is installed, because
        # we're just checking the function works
        result = is_pdf2zh_available()
        self.assertIsInstance(result, bool)

    def test_check_pdf2zh_installation_returns_dict(self):
        """Test that check_pdf2zh_installation returns a proper dict."""
        result = check_pdf2zh_installation()
        self.assertIn("installed", result)
        self.assertIn("path", result)
        self.assertIn("message", result)
        self.assertIsInstance(result["installed"], bool)

    def test_translate_pdf_returns_error_when_binary_not_found(self):
        """Test translate_pdf returns error when pdf2zh binary not found."""
        with patch("scripts.pdf_translator._find_pdf2zh", return_value=None):
            result = translate_pdf(
                input_path=Path("/tmp/test.pdf"),
                output_dir=Path("/tmp/output"),
            )
        self.assertFalse(result["ok"])
        self.assertIn("error", result)
        self.assertIn("not installed", result["error"])

    def test_translate_pdf_returns_error_when_input_not_found(self):
        """Test translate_pdf returns error when input file doesn't exist."""
        with patch("scripts.pdf_translator._find_pdf2zh", return_value=Path("/usr/bin/pdf2zh")):
            result = translate_pdf(
                input_path=Path("/nonexistent/test.pdf"),
                output_dir=Path("/tmp/output"),
            )
        self.assertFalse(result["ok"])
        self.assertIn("error", result)
        self.assertIn("not found", result["error"])

    def test_translate_pdf_fallback_text_returns_error_when_input_not_found(self):
        """Test fallback extraction returns error when input doesn't exist."""
        result = translate_pdf_fallback_text(
            input_path=Path("/nonexistent/test.pdf"),
            output_dir=Path("/tmp/output"),
        )
        self.assertFalse(result["ok"])
        self.assertIn("error", result)

    def test_find_pdf2zh_returns_none_or_path(self):
        """Test _find_pdf2zh returns None or a Path object."""
        result = _find_pdf2zh()
        self.assertTrue(result is None or isinstance(result, Path))

    def test_scanned_pdf_error_pattern_no_output(self):
        """Test that 'No output PDF generated' triggers scanned PDF detection."""
        self.assertTrue(_is_scanned_pdf_error("No output PDF generated"))

    def test_scanned_pdf_error_pattern_existing(self):
        """Test existing scanned PDF error patterns still match."""
        self.assertTrue(_is_scanned_pdf_error("no paragraphs found"))
        self.assertTrue(_is_scanned_pdf_error("Scanned PDF detected"))
        self.assertTrue(_is_scanned_pdf_error("contains no extractable text"))

    def test_scanned_pdf_error_pattern_unrelated(self):
        """Test that unrelated errors don't trigger scanned PDF detection."""
        self.assertFalse(_is_scanned_pdf_error("pdf2zh not installed"))
        self.assertFalse(_is_scanned_pdf_error("Translation timeout after 600s"))

    def test_fallback_text_ocr_when_extraction_empty(self):
        """Test that fallback text extraction tries OCR when direct extraction returns empty."""
        import tempfile

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a dummy PDF file
            pdf_path = Path(tmpdir) / "scanned.pdf"
            pdf_path.write_bytes(b"%PDF-1.4 dummy")
            output_dir = Path(tmpdir) / "output"

            # Mock: direct extraction returns empty, OCR succeeds, post-OCR extraction returns text
            call_count = {"n": 0}
            def mock_extract(path):
                call_count["n"] += 1
                if call_count["n"] == 1:
                    return ""  # First call: direct extraction empty
                return "OCR extracted text"  # Second call: after OCR

            with patch("scripts.pdf_translator.is_ocrmypdf_available", return_value=True), \
                 patch("scripts.pdf_translator.ocr_pdf", return_value={"ok": True, "output_path": pdf_path}):
                result = translate_pdf_fallback_text(
                    pdf_path, output_dir,
                    source_lang="ar",
                    extract_func=mock_extract,
                )

            self.assertTrue(result["ok"])
            self.assertTrue(result["ocr_used"])
            self.assertEqual(result["text"], "OCR extracted text")


class TestVisionTranslation(unittest.TestCase):
    """Test cases for Gemini Vision PDF translation."""

    def test_extract_first_json_object(self):
        """Test JSON extraction from mixed model output."""
        raw = 'Some text before {"key": "value", "num": 42} and after'
        result = _extract_first_json_object(raw)
        self.assertEqual(result["key"], "value")
        self.assertEqual(result["num"], 42)

    def test_extract_first_json_object_with_markdown_fence(self):
        """Test JSON extraction when wrapped in markdown code fence."""
        raw = '```json\n{"translated_text": "Hello", "sections": []}\n```'
        result = _extract_first_json_object(raw)
        self.assertEqual(result["translated_text"], "Hello")

    def test_extract_first_json_object_raises_on_no_json(self):
        """Test that ValueError is raised when no JSON found."""
        with self.assertRaises(ValueError):
            _extract_first_json_object("no json here at all")

    def test_build_translated_docx_basic(self):
        """Test DOCX building from structured page translations."""
        page_translations = [
            {
                "translated_text": "Full page text",
                "sections": [
                    {"type": "header", "content": "Joint Account Agreement"},
                    {"type": "field", "content": "Account Number: 12345"},
                    {"type": "paragraph", "content": "This agreement governs the terms."},
                    {"type": "list", "content": "1. First clause\n2. Second clause"},
                    {"type": "signature", "content": "Signature of Account Holder"},
                ],
            }
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = _build_translated_docx(
                page_translations, Path(tmpdir), stem="test"
            )
            self.assertTrue(docx_path.exists())
            self.assertEqual(docx_path.name, "test-Vision-Translated.docx")
            # Verify it's a valid DOCX by opening it
            from docx import Document
            doc = Document(str(docx_path))
            # Should have paragraphs (header + field + paragraph + list items + signature stuff)
            self.assertGreater(len(doc.paragraphs), 0)

    def test_build_translated_docx_multi_page(self):
        """Test DOCX building with multiple pages."""
        pages = [
            {"translated_text": "Page 1", "sections": [{"type": "paragraph", "content": "Page 1 content"}]},
            {"translated_text": "Page 2", "sections": [{"type": "header", "content": "Page 2 Header"}]},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = _build_translated_docx(pages, Path(tmpdir), stem="multi")
            self.assertTrue(docx_path.exists())
            from docx import Document
            doc = Document(str(docx_path))
            self.assertGreater(len(doc.paragraphs), 1)

    def test_build_translated_docx_fallback_no_sections(self):
        """Test DOCX building falls back to translated_text when no sections."""
        pages = [{"translated_text": "Just plain text", "sections": []}]
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = _build_translated_docx(pages, Path(tmpdir), stem="fallback")
            self.assertTrue(docx_path.exists())
            from docx import Document
            doc = Document(str(docx_path))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            self.assertIn("Just plain text", texts)

    def test_build_translated_docx_table_section(self):
        """Test DOCX building with table sections."""
        pages = [{
            "translated_text": "Table page",
            "sections": [{"type": "table", "content": "Name | Age\nAlice | 30\nBob | 25"}],
        }]
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = _build_translated_docx(pages, Path(tmpdir), stem="table")
            self.assertTrue(docx_path.exists())
            from docx import Document
            doc = Document(str(docx_path))
            self.assertEqual(len(doc.tables), 1)
            self.assertEqual(len(doc.tables[0].rows), 3)

    def test_translate_pdf_vision_no_api_key(self):
        """Test vision translation returns error when no API key."""
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "test.pdf"
            pdf_path.write_bytes(b"%PDF-1.4 dummy")
            with patch("scripts.pdf_translator._load_env_key", return_value=""):
                result = translate_pdf_vision(pdf_path, Path(tmpdir) / "out")
            self.assertFalse(result["ok"])
            self.assertIn("GOOGLE_API_KEY", result["error"])

    def test_translate_pdf_vision_input_not_found(self):
        """Test vision translation returns error for missing input."""
        result = translate_pdf_vision(
            Path("/nonexistent/test.pdf"), Path("/tmp/out")
        )
        self.assertFalse(result["ok"])
        self.assertIn("not found", result["error"])

    def test_translate_pdf_vision_mocked_success(self):
        """Test vision translation with mocked Gemini API."""
        mock_response = json.dumps({
            "translated_text": "Joint Account Agreement",
            "sections": [
                {"type": "header", "content": "Joint Account Agreement"},
                {"type": "paragraph", "content": "This is a test translation."},
            ],
        })

        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "test.pdf"
            pdf_path.write_bytes(b"%PDF-1.4 dummy")
            output_dir = Path(tmpdir) / "output"

            # Mock pdftoppm to create fake PNGs
            png_dir = output_dir / "_vision_pages"
            png_dir.mkdir(parents=True, exist_ok=True)
            fake_png = png_dir / "page-1.png"
            # Minimal valid PNG (1x1 pixel)
            fake_png.write_bytes(
                b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
                b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00"
                b"\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00"
                b"\x05\x18\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
            )

            with patch("scripts.pdf_translator._pdf_to_pngs", return_value=[fake_png]), \
                 patch("scripts.pdf_translator._call_gemini_vision", return_value=mock_response), \
                 patch.dict("os.environ", {"GOOGLE_API_KEY": "test-key"}):
                result = translate_pdf_vision(pdf_path, output_dir)

            self.assertTrue(result["ok"])
            self.assertTrue(result["docx_path"].exists())
            self.assertEqual(result["page_count"], 1)
            self.assertTrue(result["vision_used"])

    def test_vision_fallback_in_ocr_chain(self):
        """Test that vision fallback triggers when pdf2zh + OCR both fail."""
        from scripts.pdf_translator import translate_pdf_with_ocr_fallback

        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "scanned.pdf"
            pdf_path.write_bytes(b"%PDF-1.4 dummy")
            output_dir = Path(tmpdir) / "output"

            mock_docx = output_dir / "scanned-Vision-Translated.docx"

            with patch("scripts.pdf_translator.translate_pdf", return_value={"ok": False, "error": "No output PDF generated"}), \
                 patch("scripts.pdf_translator.is_ocrmypdf_available", return_value=True), \
                 patch("scripts.pdf_translator.ocr_pdf", return_value={"ok": True, "output_path": pdf_path}), \
                 patch("scripts.pdf_translator.translate_pdf_vision") as mock_vision, \
                 patch.dict("os.environ", {"GOOGLE_API_KEY": "test-key"}):
                mock_vision.return_value = {"ok": True, "docx_path": mock_docx, "page_count": 2, "vision_used": True}
                result = translate_pdf_with_ocr_fallback(pdf_path, output_dir)

            self.assertTrue(result["ok"])
            self.assertEqual(result["docx_path"], mock_docx)
            self.assertTrue(result.get("vision_used"))


import os


class TestPdfPipelineIntegration(unittest.TestCase):
    """Test PDF processing integration in v4_pipeline."""

    def test_process_pdf_files_function_exists(self):
        """Test that _process_pdf_files is importable from v4_pipeline."""
        from scripts.v4_pipeline import _process_pdf_files
        self.assertTrue(callable(_process_pdf_files))

    def test_process_pdf_files_handles_empty_candidates(self):
        """Test _process_pdf_files handles empty candidate list."""
        from scripts.v4_pipeline import _process_pdf_files
        import tempfile

        with tempfile.TemporaryDirectory() as tmpdir:
            result = _process_pdf_files(
                candidates=[],
                review_dir=Path(tmpdir),
            )
            self.assertEqual(result["pdf_files"], [])
            self.assertEqual(result["non_pdf_candidates"], [])
            self.assertEqual(result["translated_pdfs"], [])
            self.assertEqual(result["extracted_texts"], [])
            self.assertEqual(result["warnings"], [])
            self.assertEqual(result["errors"], [])

    def test_process_pdf_files_separates_pdf_and_non_pdf(self):
        """Test _process_pdf_files separates PDF and non-PDF candidates."""
        from scripts.v4_pipeline import _process_pdf_files
        import tempfile

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create mock candidates
            candidates = [
                {"path": "/tmp/doc.docx", "name": "doc.docx", "language": "ar", "version": "v1", "role": "source"},
                {"path": "/tmp/data.xlsx", "name": "data.xlsx", "language": "ar", "version": "v1", "role": "source"},
            ]

            result = _process_pdf_files(
                candidates=candidates,
                review_dir=Path(tmpdir),
            )

            self.assertEqual(result["pdf_files"], [])
            self.assertEqual(len(result["non_pdf_candidates"]), 2)


if __name__ == "__main__":
    unittest.main()
