#!/usr/bin/env python3

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from scripts.openclaw_artifact_writer import write_artifacts


def _make_docx(path: Path, text: str) -> None:
    doc = Document()
    doc.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def _make_xlsx(path: Path, text: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = text
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


class OpenClawArtifactWriterTest(unittest.TestCase):
    def test_write_artifacts_outputs_verify_bundle(self):
        with tempfile.TemporaryDirectory() as tmp:
            review = Path(tmp) / "_VERIFY" / "job_1"
            template = Path(tmp) / "english_v1.docx"
            _make_docx(template, "Baseline text")

            manifest = write_artifacts(
                review_dir=str(review),
                draft_a_template_path=str(template),
                delta_pack={"added": [], "removed": [], "modified": [], "summary_by_section": [], "stats": {}},
                model_scores={"judge_margin": 0.1, "term_hit": 0.95},
                quality={"judge_margin": 0.1, "term_hit": 0.95, "expansion_used": False},
                quality_report={"rounds": [], "convergence_reached": True, "stop_reason": "double_pass"},
                job_id="job_1",
                task_type="REVISION_UPDATE",
                confidence=0.93,
                estimated_minutes=12,
                runtime_timeout_minutes=16,
                iteration_count=2,
                double_pass=True,
                status_flags=[],
                candidate_files=[],
                review_questions=["Check table numbering."],
                draft_payload={
                    "final_text": "Final line 1",
                    "final_reflow_text": "Final reflow line 1",
                    "review_brief_points": ["Focus on Arabic V2 additions."],
                    "change_log_points": ["Updated Section 3 wording."],
                },
                plan_payload={"intent": {"task_type": "REVISION_UPDATE"}, "plan": {"estimated_minutes": 12}},
            )

            self.assertTrue(Path(manifest["final_docx"]).exists())
            self.assertTrue(Path(manifest["final_reflow_docx"]).exists())
            self.assertTrue(Path(manifest["review_brief_docx"]).exists())
            self.assertTrue(Path(manifest["change_log_md"]).exists())
            self.assertTrue(Path(manifest["execution_plan_json"]).exists())
            self.assertTrue(Path(manifest["quality_report_json"]).exists())

    def test_write_artifacts_can_emit_final_xlsx(self):
        with tempfile.TemporaryDirectory() as tmp:
            review = Path(tmp) / "_VERIFY" / "job_xlsx"
            template = Path(tmp) / "english_v1.docx"
            _make_docx(template, "Baseline text")

            manifest = write_artifacts(
                review_dir=str(review),
                draft_a_template_path=str(template),
                delta_pack={"added": [], "removed": [], "modified": [], "summary_by_section": [], "stats": {}},
                model_scores={"judge_margin": 0.12, "term_hit": 0.97},
                quality={"judge_margin": 0.12, "term_hit": 0.97, "expansion_used": False},
                quality_report={"rounds": [], "convergence_reached": True, "stop_reason": "double_pass"},
                job_id="job_xlsx",
                task_type="SPREADSHEET_TRANSLATION",
                confidence=0.9,
                estimated_minutes=8,
                runtime_timeout_minutes=10,
                iteration_count=1,
                double_pass=True,
                status_flags=[],
                candidate_files=[],
                review_questions=[],
                draft_payload={
                    "final_text": "F1\nF2",
                    "final_reflow_text": "R1",
                    "review_brief_points": [],
                    "change_log_points": ["normalized sheet header"],
                },
                generate_final_xlsx=True,
                plan_payload={"intent": {"task_type": "SPREADSHEET_TRANSLATION"}, "plan": {"estimated_minutes": 8}},
            )

            self.assertIn("final_xlsx", manifest)
            self.assertTrue(Path(manifest["final_xlsx"]).exists())

    def test_write_artifacts_emits_per_source_translated_xlsx_when_multiple_sources(self):
        with tempfile.TemporaryDirectory() as tmp:
            review = Path(tmp) / "_VERIFY" / "job_multi_xlsx"
            template = Path(tmp) / "template.docx"
            _make_docx(template, "Baseline text")

            src1 = Path(tmp) / "one.xlsx"
            src2 = Path(tmp) / "two.xlsx"
            _make_xlsx(src1, "Hello")
            _make_xlsx(src2, "Hello 2")

            manifest = write_artifacts(
                review_dir=str(review),
                draft_a_template_path=str(template),
                delta_pack={"added": [], "removed": [], "modified": [], "summary_by_section": [], "stats": {}},
                model_scores={"judge_margin": 0.12, "term_hit": 0.97},
                quality={"judge_margin": 0.12, "term_hit": 0.97, "expansion_used": False},
                quality_report={"rounds": [], "convergence_reached": True, "stop_reason": "double_pass"},
                job_id="job_multi_xlsx",
                task_type="SPREADSHEET_TRANSLATION",
                confidence=0.9,
                estimated_minutes=8,
                runtime_timeout_minutes=10,
                iteration_count=1,
                double_pass=True,
                status_flags=[],
                candidate_files=[
                    {"path": str(src1), "name": src1.name},
                    {"path": str(src2), "name": src2.name},
                ],
                review_questions=[],
                draft_payload={
                    "final_text": "ignored",
                    "final_reflow_text": "ignored",
                    "review_brief_points": [],
                    "change_log_points": [],
                    "xlsx_translation_map": [
                        {"file": src1.name, "sheet": "Sheet1", "cell": "A1", "text": "Bonjour"},
                        {"file": src2.name, "sheet": "Sheet1", "cell": "A1", "text": "Salut"},
                    ],
                },
                generate_final_xlsx=True,
                plan_payload={"intent": {"task_type": "SPREADSHEET_TRANSLATION"}, "plan": {"estimated_minutes": 8}},
            )

            self.assertIn("xlsx_files", manifest)
            self.assertEqual(len(manifest["xlsx_files"]), 2)
            for entry in manifest["xlsx_files"]:
                self.assertTrue(Path(entry["path"]).exists())
            self.assertNotIn("final_xlsx", manifest)

    def test_write_artifacts_emits_per_source_translated_docx_when_multiple_sources(self):
        with tempfile.TemporaryDirectory() as tmp:
            review = Path(tmp) / "_VERIFY" / "job_multi_docx"
            src1 = Path(tmp) / "one.docx"
            src2 = Path(tmp) / "two.docx"
            _make_docx(src1, "مرحبا")
            _make_docx(src2, "شكرا")

            manifest = write_artifacts(
                review_dir=str(review),
                draft_a_template_path=str(src1),
                delta_pack={"added": [], "removed": [], "modified": [], "summary_by_section": [], "stats": {}},
                model_scores={"judge_margin": 0.12, "term_hit": 0.97},
                quality={"judge_margin": 0.12, "term_hit": 0.97, "expansion_used": False},
                quality_report={"rounds": [], "convergence_reached": True, "stop_reason": "double_pass"},
                job_id="job_multi_docx",
                task_type="MULTI_FILE_BATCH",
                confidence=0.9,
                estimated_minutes=8,
                runtime_timeout_minutes=10,
                iteration_count=1,
                double_pass=True,
                status_flags=[],
                candidate_files=[
                    {"path": str(src1), "name": src1.name},
                    {"path": str(src2), "name": src2.name},
                ],
                review_questions=[],
                draft_payload={
                    "final_text": "ignored",
                    "final_reflow_text": "ignored",
                    "review_brief_points": [],
                    "change_log_points": [],
                    "docx_translation_map": [
                        {"file": src1.name, "id": "p:1", "text": "Hello"},
                        {"file": src2.name, "id": "p:1", "text": "Thanks"},
                    ],
                },
                plan_payload={"intent": {"task_type": "MULTI_FILE_BATCH"}, "plan": {"estimated_minutes": 8}},
            )

            self.assertIn("docx_files", manifest)
            self.assertEqual(len(manifest["docx_files"]), 2)
            self.assertEqual(manifest["primary_docx"], manifest["final_docx"])
            for entry in manifest["docx_files"]:
                self.assertTrue(Path(entry["path"]).exists())

            translated_one = review / "one_translated.docx"
            translated_two = review / "two_translated.docx"
            self.assertTrue(translated_one.exists())
            self.assertTrue(translated_two.exists())
            self.assertIn("Hello", Document(str(translated_one)).paragraphs[0].text)
            self.assertIn("Thanks", Document(str(translated_two)).paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
