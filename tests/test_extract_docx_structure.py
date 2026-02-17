#!/usr/bin/env python3
"""Tests for extract_docx_structure module enhancements."""

import json
import tempfile
import unittest
from pathlib import Path

# Import the module under test
from scripts.extract_docx_structure import extract_structure, normalize_text, has_arabic


class TestNormalizeText(unittest.TestCase):
    """Tests for normalize_text function."""

    def test_removes_nbsp(self):
        """Replace non-breaking spaces with regular spaces."""
        self.assertEqual(normalize_text("Hello\u00a0World"), "Hello World")

    def test_collapses_whitespace(self):
        """Collapse multiple spaces into single space."""
        self.assertEqual(normalize_text("Hello   World"), "Hello World")

    def test_strips_whitespace(self):
        """Strip leading and trailing whitespace."""
        self.assertEqual(normalize_text("  Hello World  "), "Hello World")


class TestHasArabic(unittest.TestCase):
    """Tests for has_arabic function."""

    def test_detects_arabic(self):
        """Detect Arabic characters."""
        self.assertTrue(has_arabic("مرحبا بالعالم"))
        self.assertTrue(has_arabic("Hello مرحبا World"))

    def test_no_arabic(self):
        """Return False for non-Arabic text."""
        self.assertFalse(has_arabic("Hello World"))
        self.assertFalse(has_arabic("Bonjour le monde"))

    def test_empty_text(self):
        """Handle empty text."""
        self.assertFalse(has_arabic(""))


class TestExtractStructureChecksums(unittest.TestCase):
    """Tests for checksum generation in extract_structure."""

    def test_checksums_present(self):
        """Verify checksums are present in output."""
        # Create a minimal docx file for testing
        # We use python-docx to create a test document
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a simple document
            doc = Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("Second paragraph")

            test_path = Path(tmpdir) / "test.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            # Verify checksums field exists
            self.assertIn("checksums", result)
            checksums = result["checksums"]

            # Verify all required checksum fields
            self.assertIn("content_checksum", checksums)
            self.assertIn("structure_checksum", checksums)
            self.assertIn("paragraph_count", checksums)
            self.assertIn("table_count", checksums)
            self.assertIn("block_count", checksums)
            self.assertIn("question_count", checksums)

            # Verify types
            self.assertEqual(len(checksums["content_checksum"]), 64)  # SHA-256 hex
            self.assertEqual(len(checksums["structure_checksum"]), 16)  # Truncated
            self.assertEqual(checksums["paragraph_count"], 2)
            self.assertEqual(checksums["table_count"], 0)
            self.assertEqual(checksums["question_count"], 0)

    def test_block_checksums_present(self):
        """Verify each block has a checksum."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")

            test_path = Path(tmpdir) / "test.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            for block in result["blocks"]:
                self.assertIn("checksum", block)
                self.assertEqual(len(block["checksum"]), 16)

    def test_table_block_checksum(self):
        """Verify table blocks have checksums."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(1, 0).text = "C"
            table.cell(1, 1).text = "D"

            test_path = Path(tmpdir) / "test.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            # Find the table block
            table_blocks = [b for b in result["blocks"] if b["kind"] == "table"]
            self.assertEqual(len(table_blocks), 1)

            table_block = table_blocks[0]
            self.assertIn("checksum", table_block)
            self.assertEqual(len(table_block["checksum"]), 16)


class TestQuestionnaireDetection(unittest.TestCase):
    """Tests for questionnaire detection in extract_structure."""

    def test_simple_questionnaire_detected(self):
        """Detect a simple questionnaire table."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()

            # Create a questionnaire table
            table = doc.add_table(rows=3, cols=6)
            # Header row
            table.cell(0, 0).text = "Question"
            table.cell(0, 1).text = "1"
            table.cell(0, 2).text = "2"
            table.cell(0, 3).text = "3"
            table.cell(0, 4).text = "4"
            table.cell(0, 5).text = "5"
            # Questions
            table.cell(1, 0).text = "How often do you code?"
            table.cell(2, 0).text = "Do you enjoy debugging?"

            test_path = Path(tmpdir) / "questionnaire.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            # Should detect questionnaire
            self.assertIn("questionnaire_info", result)
            q_info = result["questionnaire_info"]
            self.assertTrue(q_info["is_questionnaire"])
            self.assertEqual(q_info["total_questions"], 2)

            # Checksums should reflect question count
            self.assertEqual(result["checksums"]["question_count"], 2)

    def test_non_questionnaire_no_info(self):
        """Non-questionnaire documents should not have questionnaire_info."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("This is just a regular document.")
            doc.add_paragraph("No questionnaires here.")

            test_path = Path(tmpdir) / "regular.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            # Should not have questionnaire_info
            self.assertNotIn("questionnaire_info", result)
            self.assertEqual(result["checksums"]["question_count"], 0)

    def test_development_scale_questionnaire(self):
        """Detect questionnaire with Not Yet/Emerging/etc scale."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()

            table = doc.add_table(rows=3, cols=6)
            table.cell(0, 0).text = "Competency"
            table.cell(0, 1).text = "Not Yet"
            table.cell(0, 2).text = "Emerging"
            table.cell(0, 3).text = "Developing"
            table.cell(0, 4).text = "Proficient"
            table.cell(0, 5).text = "Leading"
            table.cell(1, 0).text = "Demonstrates AI knowledge"
            table.cell(2, 0).text = "Uses AI tools effectively"

            test_path = Path(tmpdir) / "competency.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)

            self.assertIn("questionnaire_info", result)
            self.assertTrue(result["questionnaire_info"]["is_questionnaire"])
            self.assertEqual(result["questionnaire_info"]["total_questions"], 2)


class TestStructureStability(unittest.TestCase):
    """Tests for checksum stability and uniqueness."""

    def test_same_content_same_checksum(self):
        """Same content produces same checksum."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create first document
            doc1 = Document()
            doc1.add_paragraph("Test paragraph")
            path1 = Path(tmpdir) / "doc1.docx"
            doc1.save(str(path1))

            # Create second document with same content
            doc2 = Document()
            doc2.add_paragraph("Test paragraph")
            path2 = Path(tmpdir) / "doc2.docx"
            doc2.save(str(path2))

            result1 = extract_structure(path1)
            result2 = extract_structure(path2)

            # Content checksums should be identical
            self.assertEqual(
                result1["checksums"]["content_checksum"],
                result2["checksums"]["content_checksum"],
            )

    def test_different_content_different_checksum(self):
        """Different content produces different checksum."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc1 = Document()
            doc1.add_paragraph("First document")
            path1 = Path(tmpdir) / "doc1.docx"
            doc1.save(str(path1))

            doc2 = Document()
            doc2.add_paragraph("Second document")
            path2 = Path(tmpdir) / "doc2.docx"
            doc2.save(str(path2))

            result1 = extract_structure(path1)
            result2 = extract_structure(path2)

            # Content checksums should be different
            self.assertNotEqual(
                result1["checksums"]["content_checksum"],
                result2["checksums"]["content_checksum"],
            )


class TestLanguageDetection(unittest.TestCase):
    """Tests for language_hint in extract_structure."""

    def test_arabic_detection(self):
        """Detect Arabic language."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("مرحبا بالعالم")

            test_path = Path(tmpdir) / "arabic.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)
            self.assertEqual(result["language_hint"], "ar")

    def test_english_detection(self):
        """Detect English language."""
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            doc = Document()
            doc.add_paragraph("Hello World")

            test_path = Path(tmpdir) / "english.docx"
            doc.save(str(test_path))

            result = extract_structure(test_path)
            self.assertEqual(result["language_hint"], "en")


if __name__ == "__main__":
    unittest.main()
